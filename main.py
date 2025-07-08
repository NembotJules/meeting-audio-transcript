import streamlit as st
from mistralai import Mistral
import os
import tempfile
from datetime import datetime, timedelta
import requests
from transformers import pipeline
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import warnings
import torch
import torchaudio
import json
import re
import base64

# Try to import tiktoken for accurate token counting; fall back to simple method if unavailable
try:
    import tiktoken
    TOKENIZER_AVAILABLE = True
except ImportError:
    TOKENIZER_AVAILABLE = False

# Suppress warnings for cleaner output
warnings.filterwarnings("ignore")

st.set_page_config(page_title="Meeting Transcription Tool", page_icon=":microphone:", layout="wide")

def transcribe_with_elevenlabs(video_file, elevenlabs_api_key):
    """Transcribe the uploaded video file using ElevenLabs API."""
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{video_file.name.split('.')[-1]}") as temp_video:
            temp_video.write(video_file.getvalue())
            temp_video_path = temp_video.name
        
        try:
            # Create transcript using ElevenLabs API
            response = requests.post(
                "https://api.elevenlabs.io/v1/speech-to-text?enable_logging=true",
                headers={
                    "xi-api-key": elevenlabs_api_key
                },
                data={
                    'model_id': "scribe_v1",
                },
                files={
                    'file': (video_file.name, open(temp_video_path, 'rb')),
                },
            )
            
            if response.status_code == 200:
                transcription = response.json()['text']
                return transcription
            else:
                st.error(f"ElevenLabs API error: {response.status_code} - {response.text}")
                return None
                
        finally:
            os.unlink(temp_video_path)
    except Exception as e:
        st.error(f"Error during ElevenLabs transcription: {e}")
        return None

def combine_transcripts_with_speakers(accurate_transcript, teams_transcript, mistral_api_key):
    """
    Combine accurate transcription with Teams transcript to identify speakers.
    Uses Mistral API to match speakers from Teams transcript with accurate content.
    """
    if not accurate_transcript or not teams_transcript or not mistral_api_key:
        return None
    
    try:
        client = Mistral(api_key=mistral_api_key)
        
        prompt = f"""
        You are an expert at matching speaker names with their spoken content.
        
        I have two transcripts of the same meeting:
        
        1. ACCURATE TRANSCRIPT (from ElevenLabs) - High quality but no speaker names:
        {accurate_transcript}
        
        2. TEAMS TRANSCRIPT - Lower quality but has speaker names:
        {teams_transcript}
        
        Your task is to create a FINAL TRANSCRIPT that combines the accuracy of the first transcript with the speaker identification from the second transcript.
        
        Instructions:
        1. Use the ACCURATE TRANSCRIPT as the base for content
        2. Identify speakers from the TEAMS TRANSCRIPT and match them to the corresponding content
        3. Format the output as: "Speaker Name: [accurate content]"
        4. Maintain the chronological order
        5. If you can't identify a speaker for a segment, use "Unknown: [content]"
        6. Preserve all important details, numbers, dates, and technical terms from the accurate transcript
        
        Return ONLY the final transcript with speaker names, no additional text or explanations.
        """
        
        response = client.chat.complete(
            model="mistral-large-latest",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.1,
            max_tokens=8000
        )
        
        final_transcript = response.choices[0].message.content.strip()
        return final_transcript
        
    except Exception as e:
        st.error(f"Error combining transcripts: {e}")
        return None

def transcribe_audio(audio_file, file_extension, model_size="base"):
    """Transcribe the uploaded audio file to text using the Whisper model"""
    try:
        model_id_mapping = {
            "tiny": "openai/whisper-tiny",
            "base": "openai/whisper-base",
            "small": "openai/whisper-small",
            "medium": "openai/whisper-medium",
            "large": "openai/whisper-large-v3",
        }
        model_id = model_id_mapping.get(model_size, "openai/whisper-base")
        transcriber = pipeline("automatic-speech-recognition", model=model_id)
        
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_extension}") as temp_audio:
            temp_audio.write(audio_file.getvalue())
            temp_audio_path = temp_audio.name
        
        try:
            waveform, sample_rate = torchaudio.load(temp_audio_path, backend="ffmpeg")
            if sample_rate != 16000:
                waveform = torchaudio.functional.resample(waveform, sample_rate, 16000)
            if waveform.shape[0] > 1:
                waveform = torch.mean(waveform, dim=0, keepdim=True)
            result = transcriber({"raw": waveform[0].numpy(), "sampling_rate": 16000}, chunk_length_s=30, stride_length_s=5)
            return result["text"]
        finally:
            os.unlink(temp_audio_path)
    except Exception as e:
        st.error(f"Error during audio transcription: {e}")
        return f"Error during audio transcription: {e}"

def extract_context_from_report(file, mistral_api_key):
    """Extract text from the uploaded file using Mistral OCR."""
    if not file or not mistral_api_key:
        return ""
    
    file_extension = file.name.split('.')[-1].lower()
    valid_extensions = ['pdf', 'png', 'jpg', 'jpeg']
    if file_extension not in valid_extensions:
        st.error("Unsupported file type. Please upload a PDF, PNG, JPG, or JPEG.")
        return ""
    
    try:
        client = Mistral(api_key=mistral_api_key)
        
        # Upload the file to Mistral's servers
        uploaded_file = client.files.upload(
            file={
                "file_name": file.name,
                "content": file.getvalue(),
            },
            purpose="ocr"
        )
        
        # Retrieve the signed URL
        signed_url = client.files.get_signed_url(file_id=uploaded_file.id)
        
        # Set document type based on file extension
        document_type = "document_url" if file_extension == 'pdf' else "image_url"
        
        # Construct the message with the correct structure
        messages = [
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": "Extract all text from the document."
                    },
                    {
                        "type": document_type,
                        document_type: signed_url.url
                    }
                ]
            }
        ]
        
        # Call the Mistral chat API
        response = client.chat.complete(
            model="mistral-small-latest",
            messages=messages
        )
        
        # Return the extracted text
        return response.choices[0].message.content.strip()
    
    except Exception as e:
        st.error(f"Error processing file with Mistral OCR: {e}")
        return ""

def answer_question_with_context(question, context, mistral_api_key):
    """Answer a question based on the extracted context using Mistral API."""
    if not context or not question or not mistral_api_key:
        return "Please provide a question, context, and Mistral API key."
    
    prompt = f"""
    As an assistant, answer the following question based on the provided context.

    **Context**:
    {context}

    **Question**:
    {question}

    **Answer**:
    """
    
    try:
        client = Mistral(api_key=mistral_api_key)
        response = client.chat.complete(
            model="mistral-large-latest",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.1,
            max_tokens=500
        )
        
        answer = response.choices[0].message.content.strip()
        if "tableau récapitulatif des sanctions" in question.lower():
            sanctions = parse_sanctions_from_text(answer)
            if sanctions:
                st.session_state.context_sanctions = sanctions
                st.sidebar.success(f"Sanctions stored in session state!")
            else:
                st.sidebar.warning("No sanctions found in context")
        return answer
    except Exception as e:
        return f"Error: {e}"

def parse_sanctions_from_text(text):
    """Parse sanctions table from text into a list of dictionaries."""
    sanctions = []
    lines = text.split("\n")
    header_found = False
    for line in lines:
        if "NOM" in line and "RAISON" in line and "MONTANT" in line:
            header_found = True
            continue
        if header_found and line.strip():
            parts = [part.strip() for part in line.split("|")]
            if len(parts) >= 5:
                sanctions.append({
                    "name": parts[0],
                    "reason": parts[1],
                    "amount": parts[2].replace(" FCFA", "").replace("FCFA", ""),
                    "date": parts[3],
                    "status": parts[4]
                })
    return sanctions if sanctions else None

def count_tokens(text):
    """Count the number of tokens in a text string."""
    if TOKENIZER_AVAILABLE:
        encoding = tiktoken.get_encoding("cl100k_base")
        return len(encoding.encode(text))
    else:
        tokens = re.findall(r'\b\w+\b|[.,!?;]', text)
        return len(tokens)

def clean_json_response(response):
    """Clean an API response to extract valid JSON content."""
    if isinstance(response, bytes):
        try:
            response = response.decode('utf-8')
        except UnicodeDecodeError as e:
            st.error(f"Failed to decode response as UTF-8: {e}")
            return None
    if not isinstance(response, str):
        st.error(f"Response is not a string or bytes: {type(response)}")
        return None
    response = response.strip()
    if not response:
        st.error("Response is empty after stripping whitespace.")
        return None
    response = response.removeprefix('```json').removesuffix('```').strip()
    response = response.removeprefix('```').removesuffix('```').strip()
    if not (response.startswith("{") or response.startswith("[")):
        json_match = re.search(r'\{.*\}|\[.*\]', response, re.DOTALL)
        if json_match:
            response = json_match.group(0)
        else:
            st.error(f"Response does not contain valid JSON: {response}")
            return None
    return response

def extract_info_fallback(transcription, meeting_title, date, previous_context=""):
    """Fallback function to extract information using improved string parsing and regex."""
    # Define expected team members (real names from historical data)
    expected_members = [
        "Grace Divine", "Vladimir SOUA", "Gael KIAMPI", "Emmanuel TEINGA",
        "Francis KAMSU", "Jordan KAMSU-KOM", "Loïc KAMENI", "Christian DJIMELI",
        "Daniel BAYECK", "Brice DZANGUE", "Sherelle KANA", "Jules NEMBOT",
        "Nour MAHAMAT", "Franklin TANDJA", "Marcellin SEUJIP", "Divine NDE",
        "Brian ELLA ELLA", "Amelin EPOH", "Franklin YOUMBI", "Cédric DONFACK",
        "Wilfried DOPGANG", "Ismaël POUNGOUM", "Éric BEIDI", "Boris ON MAKONG",
        "Charlène GHOMSI"
    ]
    
    # Use the new structure matching historical processor with proper French defaults
    extracted_data = {
        "meeting_metadata": {
            "date": date,
            "title": meeting_title
        },
        "attendance": {
            "present": [],
            "absent": []
        },
        "agenda_items": [
            "I- Relecture du Compte Rendu",
            "II- Récapitulatif des Résolutions et des Sanctions", 
            "III- Revue d'activités",
            "IV- Faits Saillants",
            "V- Divers"
        ],
        "activities_review": [],
        "resolutions_summary": [],
        "sanctions_summary": [],
        "key_highlights": [],
        "miscellaneous": [],
        # Additional fields for document generation
        "start_time": "Non spécifié",
        "end_time": "Non spécifié",
        "rapporteur": "Non spécifié",
        "president": "Non spécifié",
        "balance_amount": "Non spécifié",
        "balance_date": date
    }

    # Extract presence list
    present_match = re.search(r"(Présents|Présent|Présentes|Présente)[:\s]*([^\n]+)", transcription, re.IGNORECASE)
    absent_match = re.search(r"(Absents|Absent|Absentes|Absente)[:\s]*([^\n]+)", transcription, re.IGNORECASE)
    
    # Debug attendance extraction
    if present_match:
        st.info(f"🔍 Found Present: '{present_match.group(2).strip()}'")
    if absent_match:
        st.info(f"🔍 Found Absent: '{absent_match.group(2).strip()}'")
    
    if present_match or absent_match:
        present = present_match.group(2).strip() if present_match else ""
        absent = absent_match.group(2).strip() if absent_match else ""
        extracted_data["attendance"]["present"] = [name.strip() for name in present.split(",") if name.strip()] if present else []
        extracted_data["attendance"]["absent"] = [name.strip() for name in absent.split(",") if name.strip()] if absent else []
        
        # Debug parsed attendance
        st.info(f"🔍 Parsed Present: {extracted_data['attendance']['present']}")
        st.info(f"🔍 Parsed Absent: {extracted_data['attendance']['absent']}")
    else:
        # Fallback method if no explicit attendance found
        names = re.findall(r"\b[A-Z][a-z]+(?: [A-Z][a-z]+)?\b", transcription)
        common_words = {"Réunion", "Projet", "Président", "Rapporteur", "Solde", "Compte", "Ordre", "Agenda"}
        names = [name for name in set(names) if name not in common_words]
        if names:
            extracted_data["attendance"]["present"] = names
            extracted_data["attendance"]["absent"] = []
            st.warning(f"🔍 Fallback attendance - found names: {names}")

    # Extract agenda items - only if explicitly mentioned, otherwise keep defaults
    agenda_match = re.search(r"(Ordre du jour|Agenda)[:\s]*([\s\S]*?)(?=\n[A-Z]+:|\Z)", transcription, re.IGNORECASE)
    if agenda_match:
        agenda_items = agenda_match.group(2).strip()
        # Only replace default if we have clear agenda items
        items = [item.strip() for item in agenda_items.split("\n") if item.strip() and len(item.strip()) > 5]
        if items and len(items) >= 3:  # Only if we have substantial agenda items
            numbered_items = []
            for idx, item in enumerate(items, 1):
                if not re.match(r"^[IVXLC]+\-", item):
                    item = f"{to_roman(idx)}- {item}"
                numbered_items.append(item)
            extracted_data["agenda_items"] = numbered_items

    # Extract start time with improved patterns
    start_time_patterns = [
        r"(?:début|commence|commencée|démarré|start)[\s\w]*?(\d{1,2}[h:]\d{2})",
        r"(?:début|commence|commencée|démarré|start)[\s\w]*?(\d{1,2}h\d{2}min)",
        r"(?:début|commence|commencée|démarré|start)[\s\w]*?(\d{1,2}h)",
        r"(?:à|vers|around)\s*(\d{1,2}[h:]\d{2})",
        r"(?:à|vers|around)\s*(\d{1,2}h)",
        r"(\d{1,2}[h:]\d{2})[\s\w]*(?:début|commence|start)",
        r"(\d{1,2}h\d{2}min)[\s\w]*(?:début|commence|start)",
        r"(\d{1,2}h)[\s\w]*(?:début|commence|start)"
    ]
    
    for pattern in start_time_patterns:
        start_time_match = re.search(pattern, transcription, re.IGNORECASE)
        if start_time_match:
            time_str = start_time_match.group(1)
            # Normalize format
            time_str = time_str.replace(":", "h").replace("min", "min")
            if not time_str.endswith("h") and not time_str.endswith("min") and "h" in time_str:
                if not time_str.endswith("min"):
                    time_str += "min"
            extracted_data["start_time"] = time_str
            break

    # Extract end time with improved patterns
    end_time_patterns = [
        r"(?:fin|terminée|terminé|ended|fini)[\s\w]*?(\d{1,2}[h:]\d{2})",
        r"(?:fin|terminée|terminé|ended|fini)[\s\w]*?(\d{1,2}h\d{2}min)",
        r"(?:fin|terminée|terminé|ended|fini)[\s\w]*?(\d{1,2}h)",
        r"(?:jusqu'à|until|vers|around)[\s\w]*?(\d{1,2}[h:]\d{2})",
        r"(?:jusqu'à|until|vers|around)[\s\w]*?(\d{1,2}h)",
        r"(\d{1,2}[h:]\d{2})[\s\w]*(?:fin|terminée|end)",
        r"(\d{1,2}h\d{2}min)[\s\w]*(?:fin|terminée|end)",
        r"(\d{1,2}h)[\s\w]*(?:fin|terminée|end)"
    ]
    
    for pattern in end_time_patterns:
        end_time_match = re.search(pattern, transcription, re.IGNORECASE)
        if end_time_match:
            time_str = end_time_match.group(1)
            # Normalize format
            time_str = time_str.replace(":", "h").replace("min", "min")
            if not time_str.endswith("h") and not time_str.endswith("min") and "h" in time_str:
                if not time_str.endswith("min"):
                    time_str += "min"
            extracted_data["end_time"] = time_str
            break

    # Extract duration and calculate end time if we have start time but no end time
    if extracted_data["start_time"] != "Non spécifié" and extracted_data["end_time"] == "Non spécifié":
        duration_patterns = [
            r"(?:durée|dure|duré|lasted|pendant)[\s\w]*?(\d{1,2}h(?:\d{1,2}min)?)",
            r"(?:durée|dure|duré|lasted|pendant)[\s\w]*?(\d{1,2}h)",
            r"(?:durée|dure|duré|lasted|pendant)[\s\w]*?(\d{1,2}min)",
            r"(?:durée|dure|duré|lasted|pendant)[\s\w]*?(\d{1,2}\s*heures?)",
            # More patterns for duration
            r"(?:pendant|during)[\s\w]*?(\d{1,2}h\d{1,2})",
            r"(?:pendant|during)[\s\w]*?(\d{1,2}\s*heure[s]?\s*et\s*\d{1,2}\s*minute[s]?)",
            r"(\d{1,2}h\d{1,2})[\s\w]*(?:durée|duration)"
        ]
        
        for pattern in duration_patterns:
            duration_match = re.search(pattern, transcription, re.IGNORECASE)
            if duration_match:
                try:
                    start_time_str = extracted_data["start_time"].replace("h", ":").replace("min", "")
                    if ":" not in start_time_str:
                        start_time_str += ":00"
                    start_time = datetime.strptime(start_time_str, "%H:%M")
                    
                    duration_str = duration_match.group(1)
                    hours = 0
                    minutes = 0
                    
                    # Better duration parsing
                    if "h" in duration_str and "min" in duration_str:
                        hours_match = re.search(r"(\d+)h", duration_str)
                        minutes_match = re.search(r"(\d+)min", duration_str)
                        if hours_match:
                            hours = int(hours_match.group(1))
                        if minutes_match:
                            minutes = int(minutes_match.group(1))
                    elif "h" in duration_str:
                        hours_match = re.search(r"(\d+)h", duration_str)
                        if hours_match:
                            hours = int(hours_match.group(1))
                    elif "min" in duration_str:
                        minutes_match = re.search(r"(\d+)min", duration_str)
                        if minutes_match:
                            minutes = int(minutes_match.group(1))
                    elif "heure" in duration_str:
                        hours_match = re.search(r"(\d+)", duration_str)
                        if hours_match:
                            hours = int(hours_match.group(1))
                    
                    duration_delta = timedelta(hours=hours, minutes=minutes)
                    end_time = start_time + duration_delta
                    extracted_data["end_time"] = end_time.strftime("%Hh%Mmin")
                    break
                except (ValueError, AttributeError) as e:
                    continue

    # Extract rapporteur and president
    rapporteur_match = re.search(r"(Rapporteur|Rapporteuse)[:\s]*([A-Z][a-z]+(?: [A-Z][a-z]+)?)", transcription, re.IGNORECASE)
    president_match = re.search(r"(Président|Présidente|Prési)[:\s]*([A-Z][a-z]+(?: [A-Z][a-z]+)?)", transcription, re.IGNORECASE)
    if rapporteur_match:
        extracted_data["rapporteur"] = rapporteur_match.group(2)
    if president_match:
        extracted_data["president"] = president_match.group(2)

    # Extract balance
    balance_match = re.search(r"(solde|compte|balance)[\s\w]*?(\d+)", transcription, re.IGNORECASE)
    if balance_match:
        extracted_data["balance_amount"] = balance_match.group(2)
    balance_date_match = re.search(r"(solde|compte|balance)[\s\w]*?(\d{2}/\d{2}/\d{4})", transcription, re.IGNORECASE)
    if balance_date_match:
        extracted_data["balance_date"] = balance_date_match.group(2)

    # Extract activities review - ensure ALL expected members have entries, but preserve multiple activities per person
    activities_section = re.search(r"(Revue des activités|Activités de la semaine)[:\s]*([\s\S]*?)(?=\n[A-Z]+:|\Z)", transcription, re.IGNORECASE)
    all_activities = []
    
    if activities_section:
        activities_text = activities_section.group(2).strip()
        activities_lines = [line.strip() for line in activities_text.split("\n") if line.strip()]
        
        for line in activities_lines:
            actor_match = re.search(r"^[A-Z][a-z]+|^([A-Z][a-z]+(?: [A-Z][a-z]+)?)[\s,]", line)
            dossier_match = re.search(r"(?:dossier|sur le dossier) ([A-Za-z0-9\s]+)", line, re.IGNORECASE)
            activities_match = re.search(r"(?:activités|menées) ([A-Za-z\s,]+)", line, re.IGNORECASE)
            results_match = re.search(r"(?:résultat|obtenu) ([A-Za-z\s,]+)", line, re.IGNORECASE)
            perspectives_match = re.search(r"(?:perspectives|prévoit de) ([A-Za-z\s,]+)", line, re.IGNORECASE)
            
            if actor_match:
                actor = actor_match.group(1) if actor_match.group(1) else actor_match.group(0)
                # Add each activity as a separate entry (don't overwrite previous activities for same actor)
                all_activities.append({
                    "actor": actor,
                    "dossier": dossier_match.group(1).strip() if dossier_match else "Non spécifié",
                    "activities": activities_match.group(1).strip() if activities_match else "Non spécifié",
                    "results": results_match.group(1).strip() if results_match else "Non spécifié",
                    "perspectives": perspectives_match.group(1).strip() if perspectives_match else "Non spécifié"
                })
    
    # Ensure ALL expected members have at least one activity entry
    existing_actors = {activity.get("actor", "") for activity in all_activities}
    for member in expected_members:
        if member not in existing_actors:
            all_activities.append({
                "actor": member,
                "dossier": "Non spécifié",
                "activities": "RAS",
                "results": "RAS", 
                "perspectives": "RAS"
            })

    extracted_data["activities_review"] = all_activities

    # Extract resolutions
    resolution_section = re.search(r"(Résolution|Resolutions)[:\s]*([\s\S]*?)(?=\n[A-Z]+:|\Z)", transcription, re.IGNORECASE)
    if resolution_section:
        resolution_text = resolution_section.group(2).strip()
        resolution_lines = [line.strip() for line in resolution_text.split("\n") if line.strip()]
        for res in resolution_lines:
            responsible_match = re.search(r"(?:par|responsable|attribué à) ([A-Z][a-z]+(?: [A-Z][a-z]+)?)", res, re.IGNORECASE)
            deadline_match = re.search(r"(?:d'ici|avant le) (\d{2}/\d{2}/\d{4})", res, re.IGNORECASE)
            responsible = responsible_match.group(1) if responsible_match else "Non spécifié"
            deadline = deadline_match.group(1) if deadline_match else "Non spécifié"
            
            # Look up historical date for this resolution
            historical_date = get_historical_resolution_date("Non spécifié", responsible, res, exclude_date=date)
            resolution_date = historical_date if historical_date else date
            
            extracted_data["resolutions_summary"].append({
                "date": resolution_date,
                "dossier": "Non spécifié",
                "resolution": res,
                "responsible": responsible,
                "deadline": deadline,
                "execution_date": "",
                "status": "En cours"
            })

    # Extract sanctions - PRIORITY: Current meeting > Context > Default
    sanction_section = re.search(r"(Sanction|Amende)[:\s]*([\s\S]*?)(?=\n[A-Z]+:|\Z)", transcription, re.IGNORECASE)
    if sanction_section:
        sanction_text = sanction_section.group(2).strip()
        sanction_lines = [line.strip() for line in sanction_text.split("\n") if line.strip()]
        for sanc in sanction_lines:
            name_match = re.search(r"^[A-Z][a-z]+|^([A-Z][a-z]+(?: [A-Z][a-z]+)?)[\s,]", sanc)
            amount_match = re.search(r"(\d+)\s*(?:FCFA|XAF)?", sanc)
            reason_match = re.search(r"(?:pour|raison) ([a-zA-Z\s]+)", sanc, re.IGNORECASE)
            name = name_match.group(1) if name_match else "Non spécifié"
            amount = amount_match.group(1) if amount_match else "0"
            reason = reason_match.group(1).strip() if reason_match else sanc
            extracted_data["sanctions_summary"].append({
                "name": name,
                "reason": reason,
                "amount": amount,
                "date": date,
                "status": "Appliquée"
            })
    
    # Use context sanctions if none found in current meeting
    if not extracted_data["sanctions_summary"] and "context_sanctions" in st.session_state:
        sanctions = st.session_state.context_sanctions
        for sanction in sanctions:
            sanction["date"] = date  # Update date to current meeting
        extracted_data["sanctions_summary"] = sanctions
        st.info(f"📋 Using stored context sanctions: {len(sanctions)} sanctions")
    
    # Default if no sanctions found anywhere
    if not extracted_data["sanctions_summary"]:
        extracted_data["sanctions_summary"] = [{
            "name": "Aucune",
            "reason": "Aucune sanction mentionnée",
            "amount": "0",
            "date": date,
            "status": "Non appliquée"
        }]

    # Apply missing data from history (excluding current meeting)
    extracted_data = smart_historical_data_filling(extracted_data, date, allow_circular=False)

    return extracted_data

def to_roman(num):
    """Convert an integer to a Roman numeral."""
    roman_numerals = {
        1: "I", 2: "II", 3: "III", 4: "IV", 5: "V",
        6: "VI", 7: "VII", 8: "VIII", 9: "IX", 10: "X"
    }
    return roman_numerals.get(num, str(num))

def load_historical_meetings(context_dir="processed_meetings", max_meetings=3, exclude_date=None):
    """Load the most recent historical meetings for context and missing data, excluding current meeting date."""
    try:
        if not os.path.exists(context_dir):
            return []
        
        # Get all JSON files in the context directory
        json_files = []
        for file in os.listdir(context_dir):
            if file.endswith('.json'):
                filepath = os.path.join(context_dir, file)
                
                # Skip files that match the current meeting date to avoid circular context
                if exclude_date:
                    # Check if filename contains the exclude_date
                    exclude_date_formatted = exclude_date.replace("/", "-")
                    if exclude_date_formatted in file:
                        st.info(f"🚫 Excluding current meeting {file} from historical context to avoid circular reference")
                        continue
                
                # Get file modification time for sorting
                mtime = os.path.getmtime(filepath)
                json_files.append((mtime, filepath, file))
        
        if not json_files:
            return []
        
        # Sort by modification time (most recent first) and take the last max_meetings
        json_files.sort(key=lambda x: x[0], reverse=True)
        recent_files = json_files[:max_meetings]
        
        meetings = []
        for mtime, filepath, filename in recent_files:
            try:
                with open(filepath, 'r', encoding='utf-8') as f:
                    meeting_data = json.load(f)
                meetings.append(meeting_data)
            except Exception as e:
                st.warning(f"Error loading {filename}: {str(e)}")
                continue
        
        return meetings
        
    except Exception as e:
        st.warning(f"Error loading historical meetings: {str(e)}")
        return []

def get_missing_data_from_history(extracted_data, historical_meetings):
    """Fill in missing data from historical meetings."""
    if not historical_meetings:
        return extracted_data
    
    # Use the most recent meeting as the primary source for missing data
    latest_meeting = historical_meetings[0]
    
    # More aggressive sanctions handling - check if sanctions are missing or just default
    current_sanctions = extracted_data.get("sanctions_summary", [])
    has_real_sanctions = any(s.get("name", "") not in ["Aucune", "", "Non spécifié"] for s in current_sanctions)
    
    if not has_real_sanctions:  # No real sanctions found in current meeting
        latest_sanctions = latest_meeting.get("sanctions_summary", [])
        historical_real_sanctions = [s for s in latest_sanctions if s.get("name", "") not in ["Aucune", "", "Non spécifié"]]
        
        if historical_real_sanctions:
            # Update dates to current meeting and copy sanctions
            updated_sanctions = []
            for sanction in historical_real_sanctions:
                updated_sanction = sanction.copy()
                updated_sanction["date"] = extracted_data["meeting_metadata"]["date"]
                updated_sanctions.append(updated_sanction)
            
            extracted_data["sanctions_summary"] = updated_sanctions
            st.info(f"📋 Pulled {len(updated_sanctions)} sanctions from previous meeting ({latest_meeting['meeting_metadata']['date']})")
            
            # Log what sanctions were pulled
            sanction_names = [s.get("name", "") for s in updated_sanctions]
            st.write(f"Sanctions carried forward: {', '.join(sanction_names)}")
    
    # If no activities, try to get perspectives from previous meeting as starting point
    if not extracted_data.get("activities_review") or len(extracted_data["activities_review"]) == 0:
        latest_activities = latest_meeting.get("activities_review", [])
        if latest_activities:
            # Convert perspectives to new activities
            new_activities = []
            for activity in latest_activities:
                perspectives = activity.get("perspectives", "")
                if perspectives and perspectives not in ["RAS", "Aucune", "", "À définir", "Non spécifié"]:
                    new_activities.append({
                        "actor": activity.get("actor", ""),
                        "dossier": activity.get("dossier", ""),
                        "activities": f"Continuation: {perspectives}",
                        "results": "À déterminer",
                        "perspectives": "À définir"
                    })
            if new_activities:
                extracted_data["activities_review"] = new_activities
                st.info(f"📋 Generated {len(new_activities)} activity continuations from previous meeting")
    
    return extracted_data

def extract_info(transcription, meeting_title, date, mistral_api_key, previous_context="", test_mode=False):
    """Extract key information from the transcription using Mistral API with historical context."""
    if not transcription or not mistral_api_key:
        return extract_info_fallback(transcription, meeting_title, date, previous_context)

    # Load historical meetings for context (allow circular reference in test mode)
    exclude_date = None if test_mode else date
    historical_meetings = load_historical_meetings(exclude_date=exclude_date)
    
    if test_mode and historical_meetings:
        st.info(f"🧪 Test mode: Loaded {len(historical_meetings)} meetings including current for perfect data completion")
    
    # Create comprehensive historical context with full JSON content
    historical_context = ""
    if historical_meetings:
        historical_context = "\n\n=== COMPLETE HISTORICAL MEETING DATA ===\n"
        historical_context += "Use this data to understand continuity, ongoing activities, pending resolutions, and current sanctions.\n"
        historical_context += "Pay special attention to sanctions_summary for ongoing sanctions that should continue.\n\n"
        
        for i, meeting in enumerate(historical_meetings[:2]):  # Use last 2 meetings for full context
            historical_context += f"=== MEETING {i+1}: {meeting['meeting_metadata']['date']} ===\n"
            
            # Include attendance for continuity
            if meeting.get("attendance"):
                historical_context += f"Attendance:\n"
                historical_context += f"- Present: {', '.join(meeting['attendance'].get('present', []))}\n"
                historical_context += f"- Absent: {', '.join(meeting['attendance'].get('absent', []))}\n\n"
            
            # Include activities with full detail
            if meeting.get("activities_review"):
                historical_context += f"Activities Review:\n"
                for activity in meeting["activities_review"]:
                    historical_context += f"- Actor: {activity.get('actor', '')}\n"
                    historical_context += f"  Dossier: {activity.get('dossier', '')}\n"
                    historical_context += f"  Activities: {activity.get('activities', '')}\n"
                    historical_context += f"  Results: {activity.get('results', '')}\n"
                    historical_context += f"  Perspectives: {activity.get('perspectives', '')}\n\n"
            
            # Include resolutions with full detail
            if meeting.get("resolutions_summary"):
                historical_context += f"Resolutions:\n"
                for resolution in meeting["resolutions_summary"]:
                    historical_context += f"- Date: {resolution.get('date', '')}\n"
                    historical_context += f"  Dossier: {resolution.get('dossier', '')}\n"
                    historical_context += f"  Resolution: {resolution.get('resolution', '')}\n"
                    historical_context += f"  Responsible: {resolution.get('responsible', '')}\n"
                    historical_context += f"  Deadline: {resolution.get('deadline', '')}\n"
                    historical_context += f"  Status: {resolution.get('status', '')}\n\n"
            
            # CRITICALLY IMPORTANT: Include sanctions with full detail
            if meeting.get("sanctions_summary"):
                historical_context += f"SANCTIONS (IMPORTANT - these may continue to next meeting):\n"
                for sanction in meeting["sanctions_summary"]:
                    if sanction.get("name", "") != "Aucune":  # Only include real sanctions
                        historical_context += f"- Name: {sanction.get('name', '')}\n"
                        historical_context += f"  Reason: {sanction.get('reason', '')}\n"
                        historical_context += f"  Amount: {sanction.get('amount', '')} FCFA\n"
                        historical_context += f"  Date: {sanction.get('date', '')}\n"
                        historical_context += f"  Status: {sanction.get('status', '')}\n\n"
            
            # Include key highlights and miscellaneous
            if meeting.get("key_highlights"):
                historical_context += f"Key Highlights: {', '.join(meeting['key_highlights'])}\n"
            if meeting.get("miscellaneous"):
                historical_context += f"Miscellaneous: {', '.join(meeting['miscellaneous'])}\n"
            
            historical_context += "\n" + "="*60 + "\n\n"

    # Create a structured prompt matching historical processor format with proper French defaults
    prompt = f"""
    You are extracting information from a meeting transcript. Use the historical context below to understand continuity and ongoing items.

    {historical_context}

    Context from uploaded document:
    {previous_context}

    Current meeting transcript:
    {transcription}

    Meeting Date: {date}
    Meeting Title: {meeting_title}

    IMPORTANT EXTRACTION RULES:
    1. AGENDA ITEMS: Unless the transcript explicitly mentions different agenda items, use these French defaults:
       - "I- Relecture du Compte Rendu"
       - "II- Récapitulatif des Résolutions et des Sanctions"
       - "III- Revue d'activités"
       - "IV- Faits Saillants"
       - "V- Divers"
    
    2. ACTIVITIES REVIEW: Create entries for ALL team members, not just those mentioned:
       - Grace Divine, Vladimir SOUA, Gael KIAMPI, Emmanuel TEINGA
       - Francis KAMSU, Jordan KAMSU-KOM, Loïc KAMENI, Christian DJIMELI
       - Daniel BAYECK, Brice DZANGUE, Sherelle KANA, Jules NEMBOT
       - Nour MAHAMAT, Franklin TANDJA, Marcellin SEUJIP, Divine NDE
       - Brian ELLA ELLA, Amelin EPOH, Franklin YOUMBI, Cédric DONFACK
       - Wilfried DOPGANG, Ismaël POUNGOUM, Éric BEIDI, Boris ON MAKONG, Charlène GHOMSI
       - If not mentioned, use "RAS" for activities, results, and perspectives
    
    3. If the current transcript mentions ongoing sanctions or doesn't specify new sanctions, use the sanctions from the historical context with updated dates.
    4. For activities, check if they continue from previous meeting perspectives.
    5. For resolutions, check if they reference previous resolutions.
    6. Maintain continuity with historical data where appropriate.

    Extract the following information and return as JSON in this EXACT structure:
    {{
        "meeting_metadata": {{
            "date": "{date}",
            "title": "{meeting_title}"
        }},
        "attendance": {{
            "present": ["list of present attendees"],
            "absent": ["list of absent attendees"]
        }},
        "agenda_items": ["I- Relecture du Compte Rendu", "II- Récapitulatif des Résolutions et des Sanctions", "III- Revue d'activités", "IV- Faits Saillants", "V- Divers"],
        "activities_review": [
            {{
                "actor": "person name",
                "dossier": "project/file name",
                "activities": "activities performed",
                "results": "results obtained",
                "perspectives": "future plans"
            }}
        ],
        "resolutions_summary": [
            {{
                "date": "{date}",
                "dossier": "project name",
                "resolution": "resolution text",
                "responsible": "person responsible",
                "deadline": "deadline date",
                "execution_date": "",
                "status": "En cours"
            }}
        ],
        "sanctions_summary": [
            {{
                "name": "person name",
                "reason": "reason for sanction",
                "amount": "amount in FCFA",
                "date": "{date}",
                "status": "status"
            }}
        ],
        "key_highlights": ["highlight 1", "highlight 2"],
        "miscellaneous": ["misc item 1", "misc item 2"]
    }}

    Additional fields for document generation:
    - start_time: meeting start time
    - end_time: meeting end time
    - rapporteur: meeting rapporteur
    - president: meeting president
    - balance_amount: account balance
    - balance_date: balance date
    """

    try:
        client = Mistral(api_key=mistral_api_key)
        response = client.chat.complete(
            model="mistral-large-latest",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.1,
            max_tokens=8000  # Increased to handle larger context
        )
        
        raw_response = response.choices[0].message.content.strip()
        cleaned_response = clean_json_response(raw_response)
        if not cleaned_response:
            st.error("Failed to clean JSON response. Falling back.")
            return extract_info_fallback(transcription, meeting_title, date, previous_context)

        extracted_data = json.loads(cleaned_response)
        if "error" in extracted_data:
            st.error(f"API error: {extracted_data['error']}. Falling back.")
            return extract_info_fallback(transcription, meeting_title, date, previous_context)

        # Fill in missing data from historical meetings (as a backup) - using smart filling
        extracted_data = smart_historical_data_filling(extracted_data, date, allow_circular=False)

        # Add meeting metadata if not present
        if "meeting_metadata" not in extracted_data:
            extracted_data["meeting_metadata"] = {"date": date, "title": meeting_title}

        # CRITICAL FIX: Extract start_time and end_time from transcription using regex patterns
        # This ensures we capture times even if the API doesn't extract them properly
        if not extracted_data.get("start_time") or extracted_data.get("start_time") == "Non spécifié":
            # Extract start time with improved patterns
            start_time_patterns = [
                # AM/PM formats first (highest priority)
                r"(\d{1,2}:\d{2}\s*AM)",  # 06:08 AM or 06:08AM
                r"(\d{1,2}:\d{2}AM)",     # 06:08AM (no space)
                r"(\d{1,2}:\d{2}\s*am)",  # lowercase am
                r"(\d{1,2}:\d{2}am)",     # lowercase am (no space)
                # French context patterns
                r"(?:début|commence|commencée|démarré|start)[\s\w]*?(\d{1,2}[h:]\d{2})",
                r"(?:début|commence|commencée|démarré|start)[\s\w]*?(\d{1,2}h\d{2}min)",
                r"(?:début|commence|commencée|démarré|start)[\s\w]*?(\d{1,2}h)",
                r"(?:à|vers|around)\s*(\d{1,2}[h:]\d{2})",
                r"(?:à|vers|around)\s*(\d{1,2}h)",
                r"(\d{1,2}[h:]\d{2})[\s\w]*(?:début|commence|start)",
                r"(\d{1,2}h\d{2}min)[\s\w]*(?:début|commence|start)",
                r"(\d{1,2}h)[\s\w]*(?:début|commence|start)",
                # Generic time patterns (fallback)
                r"(\d{1,2}:\d{2})",  # 06:08 format
                r"(\d{1,2}h\d{2})",  # 06h08 format
            ]
            
            for pattern in start_time_patterns:
                start_time_match = re.search(pattern, transcription, re.IGNORECASE)
                if start_time_match:
                    time_str = start_time_match.group(1)
                    # Use the convert_time_format function to handle all conversions
                    converted_time = convert_time_format(time_str)
                    extracted_data["start_time"] = converted_time
                    st.info(f"🕐 Extracted start time from transcription: '{time_str}' → '{converted_time}'")
                    break

        if not extracted_data.get("end_time") or extracted_data.get("end_time") == "Non spécifié":
            # Extract end time with improved patterns - find SECOND time occurrence or end-specific patterns
            end_time_patterns = [
                # End-specific context patterns first (highest priority)
                r"(?:fin|terminée|terminé|ended|fini)[\s\w]*?(\d{1,2}:\d{2}\s*[AP]M)",
                r"(?:fin|terminée|terminé|ended|fini)[\s\w]*?(\d{1,2}:\d{2}[AP]M)",
                r"(?:fin|terminée|terminé|ended|fini)[\s\w]*?(\d{1,2}[h:]\d{2})",
                r"(?:jusqu'à|until|vers|around)[\s\w]*?(\d{1,2}:\d{2}\s*[AP]M)",
                r"(?:jusqu'à|until|vers|around)[\s\w]*?(\d{1,2}[h:]\d{2})",
                r"(\d{1,2}[h:]\d{2})[\s\w]*(?:fin|terminée|end)",
                r"(\d{1,2}h\d{2}min)[\s\w]*(?:fin|terminée|end)",
                r"(\d{1,2}h)[\s\w]*(?:fin|terminée|end)"
            ]
            
            end_time_found = False
            for pattern in end_time_patterns:
                end_time_match = re.search(pattern, transcription, re.IGNORECASE)
                if end_time_match:
                    time_str = end_time_match.group(1)
                    converted_time = convert_time_format(time_str)
                    extracted_data["end_time"] = converted_time
                    st.info(f"🕕 Extracted end time from context: '{time_str}' → '{converted_time}'")
                    end_time_found = True
                    break
            
            # If no end-specific pattern found, try to find the SECOND time occurrence
            if not end_time_found:
                # Find all AM/PM times in order
                all_times_pattern = r"(\d{1,2}:\d{2}\s*(?:AM|PM|am|pm))"
                all_times = re.findall(all_times_pattern, transcription, re.IGNORECASE)
                
                if len(all_times) >= 2:
                    # Use the second time as end time
                    end_time_str = all_times[1]
                    converted_time = convert_time_format(end_time_str)
                    extracted_data["end_time"] = converted_time
                    st.info(f"🕕 Extracted end time as second occurrence: '{end_time_str}' → '{converted_time}'")
                    end_time_found = True
                elif len(all_times) == 1:
                    # Only one time found, look for non-AM/PM times as potential end time
                    all_other_times = re.findall(r"(\d{1,2}[h:]\d{2})", transcription)
                    # Filter out the start time we already found
                    start_time_raw = extracted_data.get("start_time", "").replace("min", "").replace("h", ":")
                    other_times = [t for t in all_other_times if t.replace("h", ":") != start_time_raw]
                    
                    if other_times:
                        end_time_str = other_times[0]
                        converted_time = convert_time_format(end_time_str)
                        extracted_data["end_time"] = converted_time
                        st.info(f"🕕 Extracted end time from other time formats: '{end_time_str}' → '{converted_time}'")
                        end_time_found = True

        # Extract duration and calculate end time if we have start time but no end time
        if (extracted_data.get("start_time", "Non spécifié") != "Non spécifié" and 
            extracted_data.get("end_time", "Non spécifié") == "Non spécifié"):
            duration_patterns = [
                r"(?:durée|dure|duré|lasted|pendant)[\s\w]*?(\d{1,2}h(?:\d{1,2}min)?)",
                r"(?:durée|dure|duré|lasted|pendant)[\s\w]*?(\d{1,2}h)",
                r"(?:durée|dure|duré|lasted|pendant)[\s\w]*?(\d{1,2}min)",
                r"(?:durée|dure|duré|lasted|pendant)[\s\w]*?(\d{1,2}\s*heures?)",
            ]
            
            for pattern in duration_patterns:
                duration_match = re.search(pattern, transcription, re.IGNORECASE)
                if duration_match:
                    try:
                        start_time_str = extracted_data["start_time"].replace("h", ":").replace("min", "")
                        if ":" not in start_time_str:
                            start_time_str += ":00"
                        start_time = datetime.strptime(start_time_str, "%H:%M")
                        
                        duration_str = duration_match.group(1)
                        hours = 0
                        minutes = 0
                        
                        # Parse duration
                        if "h" in duration_str and "min" in duration_str:
                            hours_match = re.search(r"(\d+)h", duration_str)
                            minutes_match = re.search(r"(\d+)min", duration_str)
                            if hours_match:
                                hours = int(hours_match.group(1))
                            if minutes_match:
                                minutes = int(minutes_match.group(1))
                        elif "h" in duration_str:
                            hours_match = re.search(r"(\d+)h", duration_str)
                            if hours_match:
                                hours = int(hours_match.group(1))
                        elif "min" in duration_str:
                            minutes_match = re.search(r"(\d+)min", duration_str)
                            if minutes_match:
                                minutes = int(minutes_match.group(1))
                        elif "heure" in duration_str:
                            hours_match = re.search(r"(\d+)", duration_str)
                            if hours_match:
                                hours = int(hours_match.group(1))
                        
                        duration_delta = timedelta(hours=hours, minutes=minutes)
                        end_time = start_time + duration_delta
                        extracted_data["end_time"] = end_time.strftime("%Hh%Mmin")
                        st.info(f"🕕 Calculated end time from duration: {extracted_data['end_time']}")
                        break
                    except (ValueError, AttributeError) as e:
                        continue

        return extracted_data

    except Exception as e:
        st.error(f"Error extracting info: {e}. Falling back.")
        return extract_info_fallback(transcription, meeting_title, date, previous_context)

def set_cell_background(cell, rgb_color):
    """Set the background color of a table cell."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), f"{rgb_color[0]:02X}{rgb_color[1]:02X}{rgb_color[2]:02X}")
    cell._element.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=0.1, bottom=0.1, left=0.1, right=0.1):
    """Set the margins of a table cell."""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin, value in zip(['top', 'bottom', 'left', 'right'], [top, bottom, left, right]):
        margin_elm = OxmlElement(f'w:{margin}')
        margin_elm.set(qn('w:w'), str(int(value * 1440)))
        margin_elm.set(qn('w:type'), 'dxa')
        tcMar.append(margin_elm)
    tcPr.append(tcMar)

def set_table_width(table, width_in_inches):
    """Set the width of the table more aggressively."""
    table.autofit = False
    table.allow_autofit = False
    table_width = Inches(width_in_inches)
    table.width = table_width
    
    # Try to set table properties more explicitly
    try:
        # Access the table properties
        tbl = table._element
        tblPr = tbl.tblPr
        
        # Set table width property
        tblW = tblPr.find(qn('w:tblW'))
        if tblW is None:
            tblW = OxmlElement('w:tblW')
            tblPr.append(tblW)
        tblW.set(qn('w:w'), str(int(width_in_inches * 1440)))  # Convert inches to twips
        tblW.set(qn('w:type'), 'dxa')
    except Exception as e:
        st.warning(f"Could not set table width properties: {e}")
    
    # Set cell properties
    for row in table.rows:
        for cell in row.cells:
            cell.width = table_width
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def set_column_widths(table, widths_in_inches):
    """Set preferred widths for each column more aggressively."""
    for i, width in enumerate(widths_in_inches):
        for row in table.rows:
            if i < len(row.cells):
                cell = row.cells[i]
                cell.width = Inches(width)
                
                # Try to set cell width properties more explicitly
                try:
                    tc = cell._element
                    tcPr = tc.get_or_add_tcPr()
                    tcW = tcPr.find(qn('w:tcW'))
                    if tcW is None:
                        tcW = OxmlElement('w:tcW')
                        tcPr.append(tcW)
                    tcW.set(qn('w:w'), str(int(width * 1440)))  # Convert inches to twips
                    tcW.set(qn('w:type'), 'dxa')
                except Exception as e:
                    # Don't show warning for every cell, just continue
                    pass

def add_styled_paragraph(doc, text, font_name="Century", font_size=12, bold=False, color=None, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """Add a styled paragraph to the document."""
    p = doc.add_paragraph(text)
    p.alignment = alignment
    run = p.runs[0]
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    return p

def add_styled_table(doc, rows, cols, headers, data, header_bg_color=(0, 0, 0), header_text_color=(255, 255, 255), alt_row_bg_color=(192, 192, 192), column_widths=None, table_width=6.5):
    """Add a styled table to the document."""
    table = doc.add_table(rows=rows, cols=cols)
    try:
        table.style = "Table Grid"
    except KeyError:
        st.warning("The 'Table Grid' style is not available. Using default style.")
    
    set_table_width(table, table_width)
    if column_widths:
        set_column_widths(table, column_widths)
    
    for j, header in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = header
        run = cell.paragraphs[0].runs[0]
        run.font.name = "Century"
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(*header_text_color)
        set_cell_background(cell, header_bg_color)
    
    for i, row_data in enumerate(data):
        row = table.rows[i + 1]
        if (i + 1) % 2 == 0:
            for cell in row.cells:
                set_cell_background(cell, alt_row_bg_color)
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = cell_text
            run = cell.paragraphs[0].runs[0]
            run.font.name = "Century"
            run.font.size = Pt(12)
    
    return table

def add_text_in_box(doc, text, bg_color=(192, 192, 192), font_size=14, box_width_in_inches=5.0):
    """Add text inside a single-cell table with a background color."""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, box_width_in_inches)
    cell = table.cell(0, 0)
    cell.text = text
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.runs[0]
    run.font.name = "Century"
    run.font.size = Pt(font_size)
    run.font.bold = True
    set_cell_background(cell, bg_color)
    set_cell_margins(cell, top=0.2, bottom=0.2, left=0.3, right=0.3)
    return table

def fill_template_and_generate_docx(extracted_info, meeting_title, meeting_date):
    """Build the Word document from scratch using the new JSON structure."""
    try:
        # Ensure complete member data and proper structure
        extracted_info = ensure_complete_member_data(extracted_info)
        
        doc = Document()

        # Handle new structure - attendance is now a dict with present/absent arrays
        attendance = extracted_info.get("attendance", {"present": [], "absent": []})
        present_attendees = attendance.get("present", [])
        absent_attendees = attendance.get("absent", [])
        
        # Get president info
        president = extracted_info.get("president", "Non spécifié")
        if president != "Non spécifié" and present_attendees:
            for i, attendee in enumerate(present_attendees):
                if attendee.lower() == president.lower():
                    present_attendees[i] = f"{attendee} (Président)"
                    break

        # Handle agenda items - should be a list now (guaranteed to be French)
        agenda_items = extracted_info.get("agenda_items", [])
        if isinstance(agenda_items, str):
            # Fallback for old format
            agenda_list = [item.strip() for item in agenda_items.split("\n") if item.strip()]
        else:
            agenda_list = agenda_items

        # Add header box
        add_text_in_box(doc, "Direction Recherches et Investissements", bg_color=(192, 192, 192), font_size=16)
        add_styled_paragraph(doc, "COMPTE RENDU DE RÉUNION", bold=True, color=RGBColor(192, 0, 0), alignment=WD_ALIGN_PARAGRAPH.CENTER)
        
        # Use meeting_metadata if available
        meeting_metadata = extracted_info.get("meeting_metadata", {})
        display_date = meeting_metadata.get("date", extracted_info.get("date", ""))
        add_styled_paragraph(doc, display_date, bold=True, color=RGBColor(192, 0, 0), alignment=WD_ALIGN_PARAGRAPH.CENTER)
        
        add_styled_paragraph(doc, f"Heure de début: {convert_time_format(extracted_info.get('start_time', 'Non spécifié'))}", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_styled_paragraph(doc, f"Heure de fin: {convert_time_format(extracted_info.get('end_time', 'Non spécifié'))}", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        
        # Debug: Check if times are being extracted and converted
        original_start = extracted_info.get('start_time', 'Non spécifié')
        original_end = extracted_info.get('end_time', 'Non spécifié')
        converted_start = convert_time_format(original_start)
        converted_end = convert_time_format(original_end)
        
        st.info(f"🕐 Start time: {original_start} → {converted_start}")
        st.info(f"🕕 End time: {original_end} → {converted_end}")
        
        rapporteur = extracted_info.get("rapporteur", "Non spécifié")
        if rapporteur != "Non spécifié":
            add_styled_paragraph(doc, f"Rapporteur: {rapporteur}", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

        # Add attendance table with proper styling
        add_styled_paragraph(doc, "◆ LISTE DE PRÉSENCE", bold=True)
        if present_attendees or absent_attendees:
            max_rows = max(len(present_attendees), len(absent_attendees)) or 1
            
            # Create table for attendance
            attendance_table = doc.add_table(rows=max_rows + 1, cols=2)
            try:
                attendance_table.style = "Table Grid"
            except KeyError:
                pass
            
            set_table_width(attendance_table, 9.0)
            set_column_widths(attendance_table, [4.5, 4.5])
            
            # Add headers with red background
            headers = [f"Présents ({len(present_attendees)})", f"Absents ({len(absent_attendees)})"]
            for j, header in enumerate(headers):
                cell = attendance_table.cell(0, j)
                cell.text = header
                run = cell.paragraphs[0].runs[0]
                run.font.name = "Century"
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                set_cell_background(cell, (200, 0, 0))  # Red background
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add data rows with alternating colors
            attendance_data = [[present_attendees[i] if i < len(present_attendees) else "", 
                              absent_attendees[i] if i < len(absent_attendees) else ""] for i in range(max_rows)]
            
            for i, row_data in enumerate(attendance_data):
                row = attendance_table.rows[i + 1]
                
                # Add alternating row colors
                if i % 2 == 0:
                    row_color = (240, 240, 240)  # Light gray
                else:
                    row_color = (255, 255, 255)  # White
                
                for j, cell_text in enumerate(row_data):
                    cell = row.cells[j]
                    cell.text = cell_text
                    run = cell.paragraphs[0].runs[0]
                    run.font.name = "Century"
                    run.font.size = Pt(11)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    set_cell_background(cell, row_color)
        else:
            add_styled_paragraph(doc, "Aucune présence spécifiée.")

        doc.add_page_break()

        # Add agenda with improved styling
        add_styled_paragraph(doc, "ORDRE DU JOUR :", bold=True, font_size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        doc.add_paragraph()  # Add spacing
        
        if agenda_list:
            for item in agenda_list:
                # Use plain paragraph without auto-numbering
                p = doc.add_paragraph()
                # Don't use List Number style to avoid double numbering
                run = p.add_run(item)
                run.font.name = "Century"
                run.font.size = Pt(12)
        else:
            # Fallback to French defaults
            french_agenda = [
                "I- Relecture du compte rendu et adoption",
                "II- Récapitulatif des résolutions et sanctions",
                "III- Revue d'activités",
                "IV- Faits saillants",
                "V- Divers"
            ]
            for item in french_agenda:
                p = doc.add_paragraph()
                # Use plain paragraph with only Roman numerals (no auto-numbering)
                run = p.add_run(item)
                run.font.name = "Century"
                run.font.size = Pt(12)

        doc.add_page_break()

        # Add section I header
        add_styled_paragraph(doc, "I. Relecture du Compte-rendu", bold=True, font_size=14, color=RGBColor(0, 0, 0))
        add_styled_paragraph(doc, "Le compte rendu précédent n'a pas été adopté et validé.", font_size=12)
        doc.add_paragraph()  # Add spacing
        
        # Add section II header  
        add_styled_paragraph(doc, "II. Récapitulatif des résolutions et sanctions", bold=True, font_size=14, color=RGBColor(0, 0, 0))
        add_styled_paragraph(doc, "Les tableaux des résolutions et des sanctions ont été examinés et mis à jour.", font_size=12)
        doc.add_paragraph()  # Add spacing

        # PAGE BREAK HERE - so Section III starts on new page
        doc.add_page_break()
        
        # Add section III header - Activity Review FIRST
        add_styled_paragraph(doc, "III. Revue d'activités", bold=True, font_size=14, color=RGBColor(0, 0, 0))
        doc.add_paragraph()  # Add spacing

        # Add activities review (organized by department with headers) - NOW FIRST TABLE
        raw_activities = extracted_info.get("activities_review", [])
        organized_activities = organize_activities_by_department(raw_activities)
        
        st.info(f"📊 Generating activity table organized by departments: {len([a for a in organized_activities if a.get('type') == 'activity'])} activities across {len([a for a in organized_activities if a.get('type') == 'header'])} departments")
        
        add_activities_table_with_departments(doc, organized_activities)
        
        doc.add_page_break()

        # Add resolutions table (AFTER activities)
        resolutions = extracted_info.get("resolutions_summary", [])
        st.info(f"📊 Generating resolutions table with {len(resolutions)} entries")
        
        add_styled_paragraph(doc, "RÉCAPITULATIF DES RÉSOLUTIONS", bold=True, color=RGBColor(192, 0, 0))
        add_resolutions_table_with_overdue_highlighting(doc, resolutions, meeting_date.strftime("%d/%m/%Y"))
        
        doc.add_page_break()

        # Add sanctions (guaranteed to have at least one entry)
        sanctions = extracted_info.get("sanctions_summary", [])
        st.info(f"📊 Generating sanctions table with {len(sanctions)} entries")
        
        add_styled_paragraph(doc, "RÉCAPITULATIF DES SANCTIONS", bold=True, color=RGBColor(192, 0, 0))
        sanctions_data = [[s.get("name", ""), s.get("reason", ""), str(s.get("amount", "")), s.get("date", ""), s.get("status", "")] for s in sanctions]
        add_styled_table(doc, len(sanctions) + 1, 5, ["NOM", "RAISON", "MONTANT (FCFA)", "DATE", "STATUT"], sanctions_data, column_widths=[3.0, 4.5, 3.0, 2.5, 3.0], table_width=16.0)

        doc.add_page_break()
        
        # Calculate balance as sum of all sanctions
        sanctions = extracted_info.get("sanctions_summary", [])
        total_sanctions = 0
        for sanction in sanctions:
            try:
                amount = sanction.get("amount", "0")
                # Handle both string and numeric amounts
                if isinstance(amount, str):
                    amount = amount.replace(" ", "").replace("FCFA", "").replace("XAF", "")
                total_sanctions += float(amount) if amount and amount != "0" else 0
            except (ValueError, TypeError):
                continue
        
        st.info(f"💰 Calculated balance from sanctions: {total_sanctions} FCFA")
        
        # Add balance with calculated amount - fix currency format and use meeting date
        meeting_date_str = extracted_info['meeting_metadata']['date']
        add_styled_paragraph(doc, f"Solde du compte de solidarité DRI (00001-00921711101-10) est de {int(total_sanctions)} XAF au {meeting_date_str}.")

        # Save to temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            doc.save(tmp.name)
            with open(tmp.name, "rb") as f:
                docx_data = f.read()
            os.unlink(tmp.name)
        return docx_data

    except Exception as e:
        st.error(f"Error generating document: {e}")
        return None

def ensure_complete_member_data(extracted_info):
    """Ensure proper data structure and formatting, but only include members who actually have activities."""
    
    # Only work with members who actually have activities - don't force all 25 members
    activities = extracted_info.get("activities_review", [])
    
    # Filter out empty or invalid activities
    valid_activities = []
    for activity in activities:
        actor = activity.get("actor", "").strip()
        if actor and actor not in ["Non spécifié", "", "RAS"]:
            # Only include if they have meaningful activities
            activities_text = activity.get("activities", "").strip()
            if activities_text and activities_text not in ["Non spécifié", "", "RAS"]:
                valid_activities.append(activity)
    
    extracted_info["activities_review"] = valid_activities
    
    # Ensure resolutions_summary has at least one entry
    if not extracted_info.get("resolutions_summary"):
        # Try to get a reasonable resolution from history first
        historical_meetings = load_historical_meetings(exclude_date=extracted_info["meeting_metadata"]["date"])
        if historical_meetings:
            # Look for an ongoing resolution to use as default
            latest_meeting = historical_meetings[0]
            historical_resolutions = latest_meeting.get("resolutions_summary", [])
            ongoing_resolutions = [r for r in historical_resolutions if r.get("status", "") != "Exécuté"]
            if ongoing_resolutions:
                # Use the first ongoing resolution as default (preserving original date)
                default_resolution = ongoing_resolutions[0].copy()
                extracted_info["resolutions_summary"] = [default_resolution]
            else:
                # Create new default with current date
                extracted_info["resolutions_summary"] = [{
                    "date": extracted_info["meeting_metadata"]["date"],
                    "dossier": "Non spécifié",
                    "resolution": "Aucune résolution spécifique mentionnée",
                    "responsible": "Non spécifié",
                    "deadline": "Non spécifié",
                    "execution_date": "",
                    "status": "En cours"
                }]
        else:
            # No history available, create new default with current date
            extracted_info["resolutions_summary"] = [{
                "date": extracted_info["meeting_metadata"]["date"],
                "dossier": "Non spécifié",
                "resolution": "Aucune résolution spécifique mentionnée",
                "responsible": "Non spécifié",
                "deadline": "Non spécifié",
                "execution_date": "",
                "status": "En cours"
            }]
    
    # Enhance resolutions with historical data for missing members
    extracted_info = enhance_resolutions_with_history(extracted_info, extracted_info["meeting_metadata"]["date"])
    
    # Ensure sanctions_summary has at least one entry
    if not extracted_info.get("sanctions_summary") or all(s.get("name") == "Aucune" for s in extracted_info.get("sanctions_summary", [])):
        # Try to load from historical context if available
        historical_meetings = load_historical_meetings(exclude_date=extracted_info["meeting_metadata"]["date"])
        sanctions_found = False
        
        if historical_meetings:
            for meeting in historical_meetings:
                sanctions = meeting.get("sanctions_summary", [])
                real_sanctions = [s for s in sanctions if s.get("name", "") not in ["Aucune", "", "Non spécifié"]]
                if real_sanctions:
                    # Update dates and carry forward sanctions
                    updated_sanctions = []
                    for sanction in real_sanctions:
                        updated_sanction = sanction.copy()
                        updated_sanction["date"] = extracted_info["meeting_metadata"]["date"]
                        updated_sanctions.append(updated_sanction)
                    extracted_info["sanctions_summary"] = updated_sanctions
                    sanctions_found = True
                    st.info(f"📋 Carried forward {len(updated_sanctions)} sanctions from historical meeting")
                    break
        
        # If no historical sanctions, keep default
        if not sanctions_found:
            extracted_info["sanctions_summary"] = [{
                "name": "Aucune",
                "reason": "Aucune sanction mentionnée", 
                "amount": "0",
                "date": extracted_info["meeting_metadata"]["date"],
                "status": "Non appliquée"
            }]
    
    # Ensure agenda items are in French
    if not extracted_info.get("agenda_items") or any("english" in item.lower() for item in extracted_info.get("agenda_items", [])):
        extracted_info["agenda_items"] = [
            "I- Relecture du Compte Rendu",
            "II- Récapitulatif des Résolutions et des Sanctions",
            "III- Revue d'activités", 
            "IV- Faits Saillants",
            "V- Divers"
        ]
    
    return extracted_info

def show_transcript_quality_tips():
    """Display transcript quality improvement tips."""
    with st.sidebar.expander("🎯 Transcript Quality Tips", expanded=False):
        st.markdown("""
        **Poor Transcript Issues:**
        - Teams transcripts are often incomplete
        - Missing speaker identification  
        - Poor French language recognition
        - Incomplete sentences and context
        
        **Better Alternatives:**
        1. **Use Whisper transcription** (built into this app)
           - Upload audio file instead of using Teams transcript
           - Choose "medium" or "large" model for better French
        
        2. **Pre-process your audio:**
           - Clear background noise
           - Ensure good microphone quality
           - Speakers should speak clearly
        
        3. **Manual cleanup:**
           - Review and edit transcripts before processing
           - Add speaker names: "Grace: ..." 
           - Fix obvious errors
        
        4. **Use structured format:**
           - Mention agenda items clearly
           - State resolutions explicitly
           - Name speakers for activities
        """)

def smart_historical_data_filling(extracted_data, date, allow_circular=False):
    """
    Intelligently fill missing data using historical meetings as memory.
    This function allows using the current meeting's JSON if it exists to test perfect generation.
    """
    try:
        # Define department structure for member identification
        department_members = {
            "DEPARTEMENT INVESTISSEMENT": ["Grace Divine", "Vladimir SOUA", "Nour MAHAMAT", "Eric BEIDI"],
            "DEPARTEMENT PROJET": ["Marcellin SEUJIP", "Franklin TANDJA"],
            "DEPARTEMENT IA": ["Emmanuel TEINGA", "Sherelle KANA", "Jules NEMBOT", "Brice DZANGUE"],
            "DEPARTEMENT INNOVATION": ["Jordan KAMSU-KOM", "Christian DJIMELI", "Daniel BAYECK", "Brian ELLA ELLA"],
            "DEPARTEMENT ETUDE": ["Gael KIAMPI", "Francis KAMSU", "Loïc KAMENI"]
        }
        
        # Flatten to get all known team members
        all_known_members = []
        for dept_members in department_members.values():
            all_known_members.extend(dept_members)
        
        # Load historical meetings (include current if allow_circular=True for testing)
        exclude_date = None if allow_circular else date
        historical_meetings = load_historical_meetings(exclude_date=exclude_date)
        
        if not historical_meetings:
            return extracted_data
        
        st.info(f"🧠 Using historical memory from {len(historical_meetings)} meetings to fill gaps...")
        
        # Use the most recent meeting as primary memory source
        latest_meeting = historical_meetings[0]
        
        # Fill missing activity data using historical perspectives as current activities
        activities = extracted_data.get("activities_review", [])
        
        # Check if we have activities but they might be incomplete (check for members who had activities before)
        existing_actors = {activity.get("actor", "") for activity in activities}
        
        # Only check for missing members who had activities in historical meetings
        historical_activities = latest_meeting.get("activities_review", [])
        historically_active_members = {activity.get("actor", "") for activity in historical_activities 
                                     if activity.get("perspectives", "") not in ["RAS", "Non spécifié", ""]}
        
        missing_members = [member for member in historically_active_members 
                          if member in all_known_members and member not in existing_actors]
        
        if missing_members:
            st.info(f"📋 Filling {len(missing_members)} missing members who had activities in historical meetings...")
            
            # Use historical data to fill missing member activities
            historical_actors = {}
            
            # Group historical activities by actor (allowing multiple activities per person)
            for activity in historical_activities:
                actor = activity.get("actor", "")
                if actor not in historical_actors:
                    historical_actors[actor] = []
                historical_actors[actor].append(activity)
            
            # Add missing members using their historical data
            for member in missing_members:
                if member in historical_actors:
                    # Add all historical activities for this member
                    for historical_activity in historical_actors[member]:
                        perspectives = historical_activity.get("perspectives", "")
                        
                        if perspectives and perspectives not in ["RAS", "Non spécifié", ""]:
                            # Use perspectives as current activities
                            activities.append({
                                "actor": member,
                                "dossier": historical_activity.get("dossier", "Non spécifié"),
                                "activities": f"Continuation: {perspectives}",
                                "results": "En cours",
                                "perspectives": "À définir selon avancement"
                            })
                            st.write(f"   • {member} ({historical_activity.get('dossier', '')}): Continued from '{perspectives[:50]}...'")
                        else:
                            # Use the last known activity only if it had meaningful content
                            last_activities = historical_activity.get("activities", "")
                            if last_activities and last_activities not in ["RAS", "Non spécifié", ""]:
                                activities.append({
                                    "actor": member,
                                    "dossier": historical_activity.get("dossier", "Non spécifié"),
                                    "activities": last_activities,
                                    "results": "RAS",
                                    "perspectives": "RAS"
                                })
            
            extracted_data["activities_review"] = activities
        
        # Fill missing resolutions using historical data
        resolutions = extracted_data.get("resolutions_summary", [])
        if len(resolutions) == 0:
            st.info("📋 Filling missing resolutions from historical memory...")
            historical_resolutions = latest_meeting.get("resolutions_summary", [])
            
            # Get ongoing resolutions (status != "Exécuté")
            ongoing_resolutions = []
            for resolution in historical_resolutions:
                if resolution.get("status", "") != "Exécuté":
                    # Keep original date - don't update to current meeting
                    updated_resolution = resolution.copy()
                    # DON'T update the date - preserve historical assignment date
                    # updated_resolution["date"] = date  # REMOVED - this was the problem!
                    ongoing_resolutions.append(updated_resolution)
            
            if ongoing_resolutions:
                extracted_data["resolutions_summary"] = ongoing_resolutions
                st.write(f"   • Carried forward {len(ongoing_resolutions)} ongoing resolutions with original dates")
            else:
                # Create default resolution entry
                extracted_data["resolutions_summary"] = [{
                    "date": date,
                    "dossier": "Suivi général",
                    "resolution": "Suivi des activités en cours selon perspectives définies",
                    "responsible": "Tous les membres",
                    "deadline": "Prochaine réunion",
                    "execution_date": "",
                    "status": "En cours"
                }]
        
        # Fill start_time and end_time if missing
        if extracted_data.get("start_time") == "Non spécifié":
            # Try to use common meeting times or historical data
            historical_start = latest_meeting.get("start_time", "Non spécifié")
            if historical_start != "Non spécifié":
                extracted_data["start_time"] = historical_start
                st.write(f"   • Using historical start time: {historical_start}")
            else:
                extracted_data["start_time"] = "09h00min"  # Default meeting start time
                st.write(f"   • Using default start time: 09h00min")
        
        if extracted_data.get("end_time") == "Non spécifié":
            historical_end = latest_meeting.get("end_time", "Non spécifié")
            if historical_end != "Non spécifié":
                extracted_data["end_time"] = historical_end
                st.write(f"   • Using historical end time: {historical_end}")
            else:
                # Calculate end time based on start time + typical meeting duration
                try:
                    start_time_str = extracted_data["start_time"].replace("h", ":").replace("min", "")
                    if ":" not in start_time_str:
                        start_time_str += ":00"
                    start_time = datetime.strptime(start_time_str, "%H:%M")
                    end_time = start_time + timedelta(hours=2)  # 2-hour default duration
                    extracted_data["end_time"] = end_time.strftime("%Hh%Mmin")
                    st.write(f"   • Calculated end time: {extracted_data['end_time']}")
                except ValueError:
                    extracted_data["end_time"] = "11h00min"  # Default end time
                    st.write(f"   • Using default end time: 11h00min")
        
        return extracted_data
        
    except Exception as e:
        st.warning(f"Error in smart historical data filling: {str(e)}")
        return extracted_data

def test_perfect_json_generation(date="16/05/2025"):
    """Test function to load existing perfect JSON and generate document."""
    try:
        import json
        import os
        
        # Find the JSON file for the date
        date_formatted = date.replace("/", "-")
        json_file = None
        
        if os.path.exists("processed_meetings"):
            for file in os.listdir("processed_meetings"):
                if date_formatted in file and file.endswith('.json'):
                    json_file = os.path.join("processed_meetings", file)
                    break
        
        if not json_file:
            st.error(f"No JSON file found for date {date}")
            return None
        
        # Load the JSON
        with open(json_file, 'r', encoding='utf-8') as f:
            extracted_info = json.load(f)
        
        st.success(f"✅ Loaded perfect JSON from {json_file}")
        
        # Ensure complete member data
        extracted_info = ensure_complete_member_data(extracted_info)
        
        # Apply smart historical filling with circular reference allowed
        extracted_info = smart_historical_data_filling(extracted_info, date, allow_circular=True)
        
        st.info(f"📊 Final data: {len(extracted_info.get('activities_review', []))} activities, {len(extracted_info.get('resolutions_summary', []))} resolutions, {len(extracted_info.get('sanctions_summary', []))} sanctions")
        
        return extracted_info
        
    except Exception as e:
        st.error(f"Error loading perfect JSON: {str(e)}")
        return None

def group_activities_by_person_and_dossier(activities_list):
    """
    Group activities by person and then by dossier, allowing multiple activities per person/dossier.
    Returns a flattened list suitable for table generation with proper grouping.
    """
    if not activities_list:
        return []
    
    # Group activities by actor
    actor_groups = {}
    for activity in activities_list:
        actor = activity.get("actor", "").strip()
        if not actor or actor == "Non spécifié":
            continue
            
        if actor not in actor_groups:
            actor_groups[actor] = {}
        
        dossier = activity.get("dossier", "Non spécifié").strip()
        if not dossier:
            dossier = "Non spécifié"
            
        if dossier not in actor_groups[actor]:
            actor_groups[actor][dossier] = []
            
        actor_groups[actor][dossier].append(activity)
    
    # Convert to flattened table format
    flattened_activities = []
    
    # Get expected team members to ensure complete coverage
    expected_members = [
        "Grace Divine", "Vladimir SOUA", "Gael KIAMPI", "Emmanuel TEINGA",
        "Francis KAMSU", "Jordan KAMSU-KOM", "Loïc KAMENI", "Christian DJIMELI",
        "Daniel BAYECK", "Brice DZANGUE", "Sherelle KANA", "Jules NEMBOT",
        "Nour MAHAMAT", "Franklin TANDJA", "Marcellin SEUJIP", "Divine NDE",
        "Brian ELLA ELLA", "Amelin EPOH", "Franklin YOUMBI", "Cédric DONFACK",
        "Wilfried DOPGANG", "Ismaël POUNGOUM", "Éric BEIDI", "Boris ON MAKONG",
        "Charlène GHOMSI"
    ]
    
    for member in expected_members:
        if member in actor_groups:
            dossiers = actor_groups[member]
            first_dossier = True
            
            for dossier_name, dossier_activities in dossiers.items():
                if len(dossier_activities) == 1:
                    # Single activity for this dossier
                    activity = dossier_activities[0]
                    flattened_activities.append({
                        "actor": member if first_dossier else "",  # Only show name on first row
                        "dossier": dossier_name,
                        "activities": activity.get("activities", "RAS"),
                        "results": activity.get("results", "RAS"),
                        "perspectives": activity.get("perspectives", "RAS")
                    })
                    first_dossier = False
                else:
                    # Multiple activities for this dossier - combine them
                    combined_activities = []
                    combined_results = []
                    combined_perspectives = []
                    
                    for activity in dossier_activities:
                        act = activity.get("activities", "").strip()
                        res = activity.get("results", "").strip() 
                        per = activity.get("perspectives", "").strip()
                        
                        if act and act != "RAS" and act != "Non spécifié":
                            combined_activities.append(f"• {act}")
                        if res and res != "RAS" and res != "Non spécifié":
                            combined_results.append(f"• {res}")
                        if per and per != "RAS" and per != "Non spécifié":
                            combined_perspectives.append(f"• {per}")
                    
                    flattened_activities.append({
                        "actor": member if first_dossier else "",  # Only show name on first row
                        "dossier": dossier_name,
                        "activities": "\n".join(combined_activities) if combined_activities else "RAS",
                        "results": "\n".join(combined_results) if combined_results else "RAS", 
                        "perspectives": "\n".join(combined_perspectives) if combined_perspectives else "RAS"
                    })
                    first_dossier = False
        else:
            # Member not found in activities - add default entry
            flattened_activities.append({
                "actor": member,
                "dossier": "Non spécifié",
                "activities": "RAS",
                "results": "RAS",
                "perspectives": "RAS"
            })
    
    return flattened_activities

def organize_activities_by_department(activities_list):
    """
    Organize activities by department as shown in the real meeting documents.
    Returns activities grouped by department with headers.
    """
    if not activities_list:
        return []
    
    # Define department structure based on the screenshots
    departments = {
        "DEPARTEMENT INVESTISSEMENT": ["Grace Divine", "Vladimir SOUA", "Nour MAHAMAT", "Eric BEIDI"],
        "DEPARTEMENT PROJET": ["Marcellin SEUJIP", "Franklin TANDJA"],
        "DEPARTEMENT IA": ["Emmanuel TEINGA", "Sherelle KANA", "Jules NEMBOT", "Brice DZANGUE"],
        "DEPARTEMENT INNOVATION": ["Jordan KAMSU-KOM", "Christian DJIMELI", "Daniel BAYECK", "Brian ELLA ELLA"],
        "DEPARTEMENT ETUDE": ["Gael KIAMPI", "Francis KAMSU", "Loïc KAMENI"]
    }
    
    # Group activities by actor first
    actor_activities = {}
    for activity in activities_list:
        actor = activity.get("actor", "").strip()
        if not actor or actor == "Non spécifié":
            continue
            
        if actor not in actor_activities:
            actor_activities[actor] = {}
        
        dossier = activity.get("dossier", "Non spécifié").strip()
        if not dossier:
            dossier = "Non spécifié"
            
        if dossier not in actor_activities[actor]:
            actor_activities[actor][dossier] = []
            
        actor_activities[actor][dossier].append(activity)
    
    # Organize by departments
    organized_activities = []
    
    for dept_name, dept_members in departments.items():
        dept_has_activities = False
        dept_activities = []
        
        for member in dept_members:
            if member in actor_activities:
                dept_has_activities = True
                dossiers = actor_activities[member]
                first_dossier = True
                
                for dossier_name, dossier_activities in dossiers.items():
                    if len(dossier_activities) == 1:
                        # Single activity for this dossier
                        activity = dossier_activities[0]
                        dept_activities.append({
                            "type": "activity",
                            "actor": member if first_dossier else "",
                            "dossier": dossier_name,
                            "activities": activity.get("activities", "RAS"),
                            "results": activity.get("results", "RAS"),
                            "perspectives": activity.get("perspectives", "RAS")
                        })
                        first_dossier = False
                    else:
                        # Multiple activities for this dossier - combine them
                        combined_activities = []
                        combined_results = []
                        combined_perspectives = []
                        
                        for activity in dossier_activities:
                            act = activity.get("activities", "").strip()
                            res = activity.get("results", "").strip() 
                            per = activity.get("perspectives", "").strip()
                            
                            if act and act != "RAS" and act != "Non spécifié":
                                combined_activities.append(f"• {act}")
                            if res and res != "RAS" and res != "Non spécifié":
                                combined_results.append(f"• {res}")
                            if per and per != "RAS" and per != "Non spécifié":
                                combined_perspectives.append(f"• {per}")
                        
                        dept_activities.append({
                            "type": "activity",
                            "actor": member if first_dossier else "",
                            "dossier": dossier_name,
                            "activities": "\n".join(combined_activities) if combined_activities else "RAS",
                            "results": "\n".join(combined_results) if combined_results else "RAS", 
                            "perspectives": "\n".join(combined_perspectives) if combined_perspectives else "RAS"
                        })
                        first_dossier = False
        
        # Add department header and activities if department has any activities
        if dept_has_activities:
            # Add department header
            organized_activities.append({
                "type": "header",
                "department": dept_name,
                "actor": "",
                "dossier": "",
                "activities": "",
                "results": "",
                "perspectives": ""
            })
            # Add department activities
            organized_activities.extend(dept_activities)
    
    return organized_activities

def add_activities_table_with_departments(doc, organized_activities, table_width=24.0):  # EXTREMELY wide table
    """Add activities table with department headers and proper styling."""
    if not organized_activities:
        return None
    
    # Define department colors - darker for headers, lighter for rows
    dept_colors = {
        "DEPARTEMENT INVESTISSEMENT": {
            "header": (180, 0, 0),      # Darker red for header
            "rows": (255, 220, 220)     # Light red for rows
        },
        "DEPARTEMENT PROJET": {
            "header": (0, 100, 0),      # Darker green for header  
            "rows": (220, 255, 220)     # Light green for rows
        },
        "DEPARTEMENT IA": {
            "header": (100, 0, 100),    # Darker purple for header
            "rows": (240, 220, 240)     # Light purple for rows
        },
        "DEPARTEMENT INNOVATION": {
            "header": (0, 80, 180),     # Darker blue for header
            "rows": (220, 240, 255)     # Light blue for rows
        },
        "DEPARTEMENT ETUDE": {
            "header": (200, 100, 0),    # Darker orange for header
            "rows": (255, 240, 220)     # Light orange for rows
        }
    }
    
    # Count total rows needed (headers + activities)
    total_rows = len(organized_activities) + 1  # +1 for main header
    
    # Create table
    table = doc.add_table(rows=total_rows, cols=5)
    try:
        table.style = "Table Grid"
    except KeyError:
        pass  # Use default style if Table Grid not available
    
    set_table_width(table, table_width)
    # EXTREMELY long/wide individual cells - make them as wide as possible:
    # [Actor, Dossier, Activities, Results, Perspectives]
    set_column_widths(table, [5.0, 5.0, 8.0, 5.0, 5.0])  # Total = 28.0 inches - MASSIVE!
    
    # Add main header row
    headers = ["ACTEURS", "DOSSIERS", "ACTIVITÉS", "RÉSULTATS", "PERSPECTIVES"]
    for j, header in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = header
        run = cell.paragraphs[0].runs[0]
        run.font.name = "Century"
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, (0, 0, 0))  # Black header
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add data rows - REWRITTEN APPROACH
    current_row = 1
    current_department = None
    
    # First pass: Create all rows with data
    row_data = []  # Store all row information for processing
    
    for item in organized_activities:
        if item.get("type") == "header":
            # Department header row
            dept_name = item.get("department", "")
            current_department = dept_name
            row_data.append({
                "type": "header",
                "department": dept_name,
                "current_department": dept_name
            })
        else:
            # Regular activity row
            actor = item.get("actor", "").strip()
            row_data.append({
                "type": "activity",
                "actor": actor,
                "dossier": item.get("dossier", ""),
                "activities": item.get("activities", ""),
                "results": item.get("results", ""),
                "perspectives": item.get("perspectives", ""),
                "current_department": current_department
            })
    
    # Second pass: Fill the table and track person spans
    person_spans = {}  # {person_name: {"start": row_num, "end": row_num, "rows": [row_nums]}}
    
    for i, item in enumerate(row_data):
        row = table.rows[current_row]
        
        if item["type"] == "header":
            # Department header row
            dept_name = item["department"]
            
            # Force header row to be very compact - multiple methods
            try:
                row.height = Inches(0.25)  # Even smaller
                row.height_rule = 1  # Exact height
            except:
                pass
            
            # Try setting height at XML level too
            try:
                tr = row._element
                trPr = tr.get_or_add_trPr()
                trHeight = OxmlElement('w:trHeight')
                trHeight.set(qn('w:val'), str(int(0.25 * 1440)))  # 0.25 inches in twips
                trHeight.set(qn('w:hRule'), 'exact')
                trPr.append(trHeight)
            except:
                pass
            
            # Merge all cells in this row for department header
            for col in range(5):
                cell = row.cells[col]
                if col == 0:
                    cell.text = dept_name
                    run = cell.paragraphs[0].runs[0]
                    run.font.name = "Century"
                    run.font.size = Pt(9)  # Even smaller font size for very compact headers
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    
                    # Very tight margins for compact headers
                    set_cell_margins(cell, top=0.02, bottom=0.02, left=0.05, right=0.05)
                else:
                    cell.text = ""
                    # Apply same very tight margins to all cells in header row
                    set_cell_margins(cell, top=0.02, bottom=0.02, left=0.05, right=0.05)
                
                # Set department header color (darker)
                dept_color_info = dept_colors.get(dept_name, {"header": (128, 128, 128), "rows": (240, 240, 240)})
                set_cell_background(cell, dept_color_info["header"])
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Now merge the cells in this row
            try:
                start_cell = row.cells[0]
                for col in range(1, 5):
                    start_cell.merge(row.cells[col])
            except Exception as e:
                st.warning(f"Could not merge header row: {e}")
            
        else:
            # Regular activity row
            actor = item["actor"]
            current_dept = item["current_department"]
            
            # Track person spans
            if actor:
                if actor not in person_spans:
                    person_spans[actor] = {"start": current_row, "end": current_row, "rows": [current_row]}
                else:
                    person_spans[actor]["end"] = current_row
                    person_spans[actor]["rows"].append(current_row)
            
            data = [
                actor,  # We'll handle the spanning later
                item["dossier"],
                item["activities"],
                item["results"],
                item["perspectives"]
            ]
            
            for j, cell_text in enumerate(data):
                cell = row.cells[j]
                cell.text = cell_text
                run = cell.paragraphs[0].runs[0]
                run.font.name = "Century"
                run.font.size = Pt(12)
                
                # Make person names bold
                if j == 0 and cell_text:  # Actor column with text
                    run.font.bold = True
                
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                # Set department row color (lighter version of department color)
                if current_dept:
                    dept_color_info = dept_colors.get(current_dept, {"header": (128, 128, 128), "rows": (240, 240, 240)})
                    set_cell_background(cell, dept_color_info["rows"])
                else:
                    # Fallback alternating colors
                    if current_row % 2 == 0:
                        set_cell_background(cell, (240, 240, 240))  # Light gray
        
        current_row += 1
    
    # Third pass: Simple name spanning - just clear duplicates (no complex merging)
    for person, span_info in person_spans.items():
        if len(span_info["rows"]) > 1:  # Person has multiple dossiers
            # Just clear duplicate names from subsequent rows - simple and reliable
            for row_num in span_info["rows"][1:]:
                try:
                    table.cell(row_num, 0).text = ""
                except:
                    pass
    
    return table

def show_meeting_guidance():
    """Display comprehensive meeting guidance for better transcription and data extraction."""
    with st.sidebar.expander("📋 Meeting Structure Guide", expanded=False):
        st.markdown("""
        ### 🎯 For Better Data Extraction
        
        **📢 Meeting Opening (Speaker should say):**
        ```
        "Bonjour, nous commençons la réunion à [HEURE]
        Présents: [List all present members clearly]
        Absents: [List absent members with reasons]
        Rapporteur: [Name]
        Président: [Name]"
        ```
        
        **📋 Activity Review Structure:**
        ```
        Department by department:
        "DEPARTEMENT INVESTISSEMENT:
        - Grace Divine: Dossier [NAME], Activités [DETAILS], 
          Résultats [DETAILS], Perspectives [DETAILS]
        - Vladimir SOUA: ..."
        ```
        
        **📜 Resolutions (Speaker should say):**
        ```
        "RESOLUTION: [Clear description]
        Responsable: [Person name]
        Échéance: [Date DD/MM/YYYY]
        Statut: En cours/Exécuté"
        ```
        
        **💰 Sanctions (If any):**
        ```
        "SANCTION: [Person] pour [Reason] 
        Montant: [Amount] FCFA
        Date: [DD/MM/YYYY]"
        ```
        
        **⏰ Meeting End:**
        ```
        "La réunion se termine à [HEURE]
        Solde du compte: [Amount] FCFA au [Date]"
        ```
        """)

    with st.sidebar.expander("🎙️ Audio Quality Tips", expanded=False):
        st.markdown("""
        ### 🔊 Recording Best Practices
        
        **Before Meeting:**
        - Use good microphone (not laptop mic)
        - Test audio levels
        - Quiet room, minimal background noise
        - Position mic centrally
        
        **During Meeting:**
        - Speak clearly and slowly
        - State your name before speaking
        - Pause between topics
        - Repeat important information
        - Spell out complex terms
        
        **For Virtual Meetings:**
        - Use "Record" feature in Teams/Zoom
        - Ask participants to mute when not speaking
        - Use headphones to reduce echo
        - Enable auto-transcription if available
        
        **Post-Meeting:**
        - Review transcript before uploading
        - Fix obvious errors manually
        - Add missing speaker names
        - Clarify unclear sections
        """)

    with st.sidebar.expander("🏗️ Meeting Template", expanded=False):
        st.markdown("""
        ### 📝 Suggested Meeting Flow
        
        **1. Opening (5 min)**
        - Time, attendance, roles
        
        **2. Previous Minutes Review (10 min)**
        - Quick validation
        
        **3. Resolutions & Sanctions Review (10 min)**
        - Update status of previous items
        
        **4. Activity Review by Department (30 min)**
        - INVESTISSEMENT → PROJET → IA → INNOVATION → ETUDE
        - Each person: Dossier, Activities, Results, Perspectives
        
        **5. Key Highlights (10 min)**
        - Important announcements
        
        **6. Miscellaneous (10 min)**
        - Other topics
        
        **7. Closing (5 min)**
        - End time, account balance
        """)

def enhance_resolutions_with_history(extracted_data, date):
    """Enhance resolutions by adding missing members' resolutions from previous meetings."""
    try:
        # Get current resolutions and find mentioned members
        current_resolutions = extracted_data.get("resolutions_summary", [])
        current_responsible = {res.get("responsible", "") for res in current_resolutions}
        
        # Expected team members
        expected_members = [
            "Grace Divine", "Vladimir SOUA", "Gael KIAMPI", "Emmanuel TEINGA",
            "Francis KAMSU", "Jordan KAMSU-KOM", "Loïc KAMENI", "Christian DJIMELI",
            "Daniel BAYECK", "Brice DZANGUE", "Sherelle KANA", "Jules NEMBOT",
            "Nour MAHAMAT", "Franklin TANDJA", "Marcellin SEUJIP", "Divine NDE",
            "Brian ELLA ELLA", "Amelin EPOH", "Franklin YOUMBI", "Cédric DONFACK",
            "Wilfried DOPGANG", "Ismaël POUNGOUM", "Éric BEIDI", "Boris ON MAKONG",
            "Charlène GHOMSI"
        ]
        
        # Find members missing from current resolutions
        missing_members = []
        for member in expected_members:
            if member not in current_responsible:
                missing_members.append(member)
        
        if missing_members:
            st.info(f"🔍 Looking for resolutions for {len(missing_members)} missing members: {', '.join(missing_members)}")
            
            # Load historical meetings
            historical_meetings = load_historical_meetings(exclude_date=date)
            
            enhanced_resolutions = list(current_resolutions)  # Start with current
            
            for missing_member in missing_members:
                # Look for this member's latest resolution in historical meetings
                for meeting in historical_meetings:
                    historical_resolutions = meeting.get("resolutions_summary", [])
                    for resolution in historical_resolutions:
                        if resolution.get("responsible", "") == missing_member:
                            # Found a resolution for this missing member
                            enhanced_resolution = resolution.copy()
                            # DON'T update the date - preserve original assignment date
                            # enhanced_resolution["date"] = date  # REMOVED - this was the problem!
                            enhanced_resolutions.append(enhanced_resolution)
                            st.info(f"   ✅ Added resolution for {missing_member} from {meeting['meeting_metadata']['date']} (original date: {resolution.get('date', 'Unknown')})")
                            break
                    else:
                        continue  # Continue to next meeting
                    break  # Found resolution, stop looking for this member
            
            extracted_data["resolutions_summary"] = enhanced_resolutions
            st.info(f"📋 Enhanced resolutions: {len(current_resolutions)} current + {len(enhanced_resolutions) - len(current_resolutions)} historical = {len(enhanced_resolutions)} total")
        
        return extracted_data
        
    except Exception as e:
        st.warning(f"Error enhancing resolutions: {str(e)}")
        return extracted_data

def get_historical_resolution_date(dossier, responsible, resolution_text, exclude_date=None):
    """
    Look up the earliest (original) assignment date for a dossier from historical meetings.
    Simply finds the earliest appearance of the dossier name across all meetings.
    
    Args:
        dossier: Dossier name to match
        responsible: Responsible person (not used - for compatibility)
        resolution_text: Resolution text (not used - for compatibility)
        exclude_date: Date to exclude from search (current meeting)
        
    Returns:
        Earliest date if found, or None if not found
    """
    try:
        # Load historical meetings
        historical_meetings = load_historical_meetings(exclude_date=exclude_date)
        
        if not historical_meetings:
            return None
        
        # If dossier is empty or "Non spécifié", can't search
        if not dossier or dossier.strip().lower() in ["non spécifié", "", "non specifie"]:
            return None
        
        # Collect all dates for this dossier (case-insensitive search)
        dossier_dates = []
        search_dossier = dossier.strip().lower()
        
        # Search through ALL historical meetings
        for meeting in historical_meetings:
            historical_resolutions = meeting.get("resolutions_summary", [])
            
            for resolution in historical_resolutions:
                hist_dossier = resolution.get("dossier", "").strip().lower()
                hist_date = resolution.get("date", "").strip()
                
                # Match if dossier names are similar (exact match or one contains the other)
                dossier_match = False
                if hist_dossier and search_dossier:
                    if hist_dossier == search_dossier:
                        dossier_match = True
                    elif len(search_dossier) > 5 and search_dossier in hist_dossier:
                        dossier_match = True
                    elif len(hist_dossier) > 5 and hist_dossier in search_dossier:
                        dossier_match = True
                
                if dossier_match and hist_date:
                    try:
                        # Parse date to ensure it's valid
                        from datetime import datetime
                        parsed_date = datetime.strptime(hist_date, "%d/%m/%Y")
                        dossier_dates.append((hist_date, parsed_date))
                        st.info(f"📅 Found dossier '{dossier}' on {hist_date}")
                    except ValueError:
                        continue
        
        if dossier_dates:
            # Sort by parsed date and return the earliest
            dossier_dates.sort(key=lambda x: x[1])
            earliest_date = dossier_dates[0][0]
            st.success(f"✅ Earliest date for dossier '{dossier}': {earliest_date}")
            return earliest_date
        else:
            st.warning(f"⚠️ No historical date found for dossier '{dossier}'")
        
        return None
        
    except Exception as e:
        st.warning(f"Error looking up historical resolution date: {str(e)}")
        return None

def add_resolutions_table_with_overdue_highlighting(doc, resolutions, current_meeting_date, table_width=20.5):
    """Add a styled resolutions table with red background for overdue items."""
    if not resolutions:
        return None
    
    # Parse current meeting date for comparison
    try:
        from datetime import datetime
        current_date = datetime.strptime(current_meeting_date, "%d/%m/%Y")
    except ValueError:
        # If current date is invalid, don't do any highlighting
        current_date = None
        st.warning(f"Invalid current meeting date format: {current_meeting_date}")
    
    # Create table
    table = doc.add_table(rows=len(resolutions) + 1, cols=8)
    try:
        table.style = "Table Grid"
    except KeyError:
        st.warning("The 'Table Grid' style is not available. Using default style.")
    
    set_table_width(table, table_width)
    column_widths = [2.2, 2.8, 4.5, 2.0, 2.5, 2.2, 2.0, 2.3]
    set_column_widths(table, column_widths)
    
    # Add headers
    headers = ["DATE", "DOSSIER", "RÉSOLUTION", "RESP.", "ÉCHÉANCE", "DATE D'EXÉCUTION", "STATUT", "COMPTE RENDU"]
    for j, header in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = header
        run = cell.paragraphs[0].runs[0]
        run.font.name = "Century"
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, (0, 0, 0))  # Black header
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add data rows with overdue highlighting
    for i, resolution in enumerate(resolutions):
        row = table.rows[i + 1]
        
        # Check if resolution is overdue
        is_overdue = False
        deadline = resolution.get("deadline", "")
        if current_date and deadline:
            try:
                deadline_date = datetime.strptime(deadline, "%d/%m/%Y")
                if deadline_date < current_date:
                    is_overdue = True
                    st.info(f"🔴 Overdue resolution: {resolution.get('dossier', '')} (deadline: {deadline})")
            except ValueError:
                # Invalid deadline format, don't highlight
                pass
        
        # Set row background color
        if is_overdue:
            row_color = (255, 200, 200)  # Light red for overdue
        elif i % 2 == 0:
            row_color = (240, 240, 240)  # Light gray for even rows
        else:
            row_color = (255, 255, 255)  # White for odd rows
        
        # Prepare row data
        row_data = [
            resolution.get("date", ""),
            resolution.get("dossier", ""),
            resolution.get("resolution", ""),
            resolution.get("responsible", ""),
            resolution.get("deadline", ""),
            resolution.get("execution_date", ""),
            resolution.get("status", ""),
            str(resolution.get("report_count", "0"))
        ]
        
        # Fill cells
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = cell_text
            run = cell.paragraphs[0].runs[0]
            run.font.name = "Century"
            run.font.size = Pt(12)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_background(cell, row_color)
    
    return table

def convert_time_format(time_str):
    """
    Convert time from various formats to French format (HHhMMmin).
    
    Args:
        time_str: Time in formats like "06:08AM", "6:08 AM", "06h08min", etc.
        
    Returns:
        Time in French format "06h08min" or original if conversion fails
    """
    if not time_str or time_str == "Non spécifié":
        return "Non spécifié"
    
    try:
        import re
        from datetime import datetime
        
        # Remove extra spaces
        time_str = time_str.strip()
        
        # If already in French format, return as is
        if re.match(r'\d{1,2}h\d{2}min?', time_str, re.IGNORECASE):
            return time_str
        
        # Handle AM/PM format: "06:08AM", "6:08 AM", etc.
        am_pm_pattern = r'(\d{1,2}):(\d{2})\s*(AM|PM)'
        match = re.match(am_pm_pattern, time_str, re.IGNORECASE)
        if match:
            hour = int(match.group(1))
            minute = int(match.group(2))
            period = match.group(3).upper()
            
            # Convert to 24-hour format
            if period == 'AM':
                if hour == 12:
                    hour = 0
            else:  # PM
                if hour != 12:
                    hour += 12
            
            return f"{hour:02d}h{minute:02d}min"
        
        # Handle 24-hour format: "14:30", "14:30:00"
        time_24h_pattern = r'(\d{1,2}):(\d{2})(?::\d{2})?'
        match = re.match(time_24h_pattern, time_str)
        if match:
            hour = int(match.group(1))
            minute = int(match.group(2))
            return f"{hour:02d}h{minute:02d}min"
        
        # If no pattern matches, return original
        return time_str
        
    except Exception as e:
        # If conversion fails, return original
        return time_str

def main():
    """
    QUICK TABLE WIDTH ADJUSTMENTS:
    
    To make activity table columns wider/narrower, edit line ~1640 in add_activities_table_with_departments():
    set_column_widths(table, [5.0, 5.0, 8.0, 5.0, 5.0])
                             ^^^^  ^^^^  ^^^^  ^^^^  ^^^^
                             Actor Doss  Act   Res   Pers
    
    - Increase numbers to make columns wider/longer horizontally
    - Decrease numbers to make columns narrower  
    - Total should be close to table_width (currently 24.0 inches)
    
    Current settings (MASSIVELY WIDE for maximum cell length):
    - Actor: 5.0 inches
    - Dossier: 5.0 inches  
    - Activities: 8.0 inches (widest)
    - Results: 5.0 inches
    - Perspectives: 5.0 inches
    - Total: 28.0 inches (MASSIVE individual cells!)
    """
    st.title("Meeting Transcription Tool")
    
    # Sidebar
    st.sidebar.header("Configuration")
    st.session_state.mistral_api_key = st.sidebar.text_input("Mistral API Key", type="password")
    st.session_state.elevenlabs_api_key = st.sidebar.text_input("ElevenLabs API Key", type="password")
    
    # Show transcript quality tips
    show_transcript_quality_tips()
    
    # Show meeting guidance for better data extraction
    show_meeting_guidance()
    
    # Show historical context information
    st.sidebar.header("📊 Historical Context")
    historical_meetings = load_historical_meetings()
    if historical_meetings:
        st.sidebar.success(f"✅ {len(historical_meetings)} historical meetings loaded")
        
        # Show recent meetings
        for i, meeting in enumerate(historical_meetings):
            date = meeting.get('meeting_metadata', {}).get('date', 'Unknown')
            st.sidebar.write(f"📅 Meeting {i+1}: {date}")
            
            # Show sanctions count
            sanctions = meeting.get('sanctions_summary', [])
            real_sanctions = [s for s in sanctions if s.get('name', '') not in ['Aucune', '', 'Non spécifié']]
            if real_sanctions:
                st.sidebar.write(f"   🚨 {len(real_sanctions)} sanctions available")
                with st.sidebar.expander(f"View sanctions from {date}"):
                    for sanction in real_sanctions[:5]:  # Show first 5
                        st.write(f"• {sanction.get('name', '')}: {sanction.get('reason', '')} ({sanction.get('amount', '')} FCFA)")
            else:
                st.sidebar.write(f"   ✅ No active sanctions")
    else:
        st.sidebar.warning("⚠️ No historical meetings found")
        st.sidebar.info("💡 Use the Historical Processor app to process past meeting documents first.")
    
    st.sidebar.header("Contexte Précédent")
    previous_report = st.sidebar.file_uploader("Télécharger le rapport précédent", type=["pdf", "png", "jpg", "jpeg"])
    if previous_report:
        st.session_state.previous_report = previous_report
        st.session_state.previous_context = ""
        st.sidebar.write("Rapport téléchargé. Posez une question pour extraire le contexte.")
    else:
        st.session_state.previous_report = None
        st.session_state.previous_context = ""
    
    # Test context
    st.sidebar.header("Tester le Contexte")
    question = st.sidebar.text_input("Posez une question sur le rapport précédent :")
    if st.sidebar.button("Poser la Question") and question:
        if not st.session_state.mistral_api_key or not st.session_state.previous_report:
            st.sidebar.error("Fournissez une clé API Mistral et un rapport précédent.")
        else:
            with st.spinner("Extraction du contexte..."):
                context = extract_context_from_report(st.session_state.previous_report, st.session_state.mistral_api_key)
                if context:
                    st.session_state.previous_context = context
                    st.sidebar.text_area("Contexte Extrait", context, height=200)
                    st.sidebar.success("Contexte extrait !")
                else:
                    st.session_state.previous_context = ""
                    st.sidebar.error("Échec de l'extraction.")
            
            with st.spinner("Obtention de la réponse..."):
                answer = answer_question_with_context(question, st.session_state.previous_context, st.session_state.mistral_api_key)
            st.sidebar.write("**Réponse :**")
            st.sidebar.write(answer)
    
    # Main content
    col1, col2 = st.columns(2)
    
    with col1:
        st.header("Détails de la RéUnion")
        meeting_title = st.text_input("Titre de la RéUnion", value="Réunion")
        meeting_date = st.date_input("Date de la RéUnion", datetime.now())
        
        # Add warning about circular reference
        formatted_date = meeting_date.strftime("%d/%m/%Y")
        st.info(f"📅 Meeting date: {formatted_date}")
        
        # Check if meeting already exists and offer testing mode
        meeting_exists = os.path.exists("processed_meetings") and any(formatted_date.replace("/", "-") in f for f in os.listdir("processed_meetings") if f.endswith('.json'))
        
        if meeting_exists:
            st.warning(f"⚠️ A meeting for {formatted_date} already exists in historical data.")
            
            # Add option for testing perfect JSON generation
            test_mode = st.checkbox(
                "🧪 **Test Mode**: Use existing JSON as perfect data source",
                help="Enable this to test document generation using the existing perfect JSON data. This allows circular reference for testing purposes."
            )
            
            if test_mode:
                st.success("✅ Test mode enabled! Will use existing JSON data to fill gaps for perfect document generation.")
                st.session_state.test_mode = True
                
                # Add test button for direct JSON loading
                if st.button("🧪 Load Perfect JSON & Generate Document"):
                    with st.spinner("Loading perfect JSON data..."):
                        perfect_data = test_perfect_json_generation(formatted_date)
                        if perfect_data:
                            st.session_state.extracted_info = perfect_data
                            st.success("✅ Perfect JSON loaded! Ready for document generation.")
                            with st.expander("📋 View Perfect JSON Data"):
                                st.json(perfect_data)
            else:
                st.info("📊 Normal mode: Will exclude existing meeting from context to avoid circular reference.")
                st.session_state.test_mode = False
        else:
            st.session_state.test_mode = False
    
    with col2:
        st.header("Transcription & Résultat")
        input_method = st.radio("Méthode d'entrée :", ("Télécharger Vidéo + Teams", "Télécharger Audio (Whisper)", "Entrer la Transcription"))
        
        if input_method == "Télécharger Vidéo + Teams":
            st.info("🎥 **Nouveau workflow:** Upload video for accurate transcription + Teams transcript for speaker names")
            
            # Video upload for ElevenLabs
            video_file = st.file_uploader("Téléchargez la vidéo de la réunion", type=["mp4", "avi", "mov", "mkv", "webm"])
            
            # Teams transcript upload
            teams_transcript = st.text_area("Collez la transcription Teams (avec noms des intervenants)", height=150, 
                                          help="Copiez-collez la transcription Teams qui contient les noms des intervenants")
            
            if video_file and teams_transcript and st.button("Traiter avec ElevenLabs + Mistral"):
                if not st.session_state.elevenlabs_api_key or not st.session_state.mistral_api_key:
                    st.error("Veuillez fournir les clés API ElevenLabs et Mistral.")
                else:
                    # Step 1: Transcribe with ElevenLabs
                    with st.spinner("Transcription avec ElevenLabs..."):
                        accurate_transcript = transcribe_with_elevenlabs(video_file, st.session_state.elevenlabs_api_key)
                        if accurate_transcript:
                            st.success("✅ Transcription ElevenLabs terminée!")
                            with st.expander("📝 Transcription précise (sans noms)"):
                                st.text_area("Transcription ElevenLabs", accurate_transcript, height=200)
                            
                            # Step 2: Combine with Teams transcript
                            with st.spinner("Combinaison avec la transcription Teams..."):
                                final_transcript = combine_transcripts_with_speakers(
                                    accurate_transcript, teams_transcript, st.session_state.mistral_api_key
                                )
                                if final_transcript:
                                    st.session_state.transcription = final_transcript
                                    st.success("✅ Transcription finale avec noms terminée!")
                                    st.text_area("Transcription finale avec noms", final_transcript, height=200)
                                    
                                    # Step 3: Extract information
                                    with st.spinner("Extraction des informations avec contexte historique..."):
                                        allow_circular = getattr(st.session_state, 'test_mode', False)
                                        if allow_circular:
                                            st.info("🧪 Test mode: Using existing JSON data for perfect generation")
                                        
                                        extracted_info = extract_info(
                                            final_transcript, meeting_title, meeting_date.strftime("%d/%m/%Y"), 
                                            st.session_state.mistral_api_key, st.session_state.get("previous_context", ""), allow_circular
                                        )
                                        if extracted_info:
                                            extracted_info = smart_historical_data_filling(extracted_info, meeting_date.strftime("%d/%m/%Y"), allow_circular=allow_circular)
                                            st.session_state.extracted_info = extracted_info
                                            st.success("✅ Information extraction completed!")
                                            with st.expander("📋 View Extracted Information"):
                                                st.json(extracted_info)
                                else:
                                    st.error("❌ Échec de la combinaison des transcriptions")
                        else:
                            st.error("❌ Échec de la transcription ElevenLabs")
        
        elif input_method == "Télécharger Audio (Whisper)":
            uploaded_file = st.file_uploader("Téléchargez un fichier audio", type=["mp3", "wav", "m4a", "flac"])
            whisper_model = st.selectbox("Modèle Whisper", ["tiny", "base", "small", "medium", "large"], index=1)
            
            st.info("💡 **Tip:** Use 'medium' or 'large' models for better French transcription quality!")
            
            if uploaded_file and st.button("Transcrire l'Audio"):
                file_extension = uploaded_file.name.split('.')[-1].lower()
                with st.spinner(f"Transcription avec Whisper {whisper_model}..."):
                    transcription = transcribe_audio(uploaded_file, file_extension, whisper_model)
                    if transcription and not transcription.startswith("Error"):
                        st.session_state.transcription = transcription
                        st.success("✅ Transcription completed!")
                        st.text_area("Transcription résultante", transcription, height=200)
                        
                        with st.spinner("Extraction des informations avec contexte historique..."):
                            allow_circular = getattr(st.session_state, 'test_mode', False)
                            if allow_circular:
                                st.info("🧪 Test mode: Using existing JSON data for perfect generation")
                            
                            extracted_info = extract_info(st.session_state.transcription, meeting_title, meeting_date.strftime("%d/%m/%Y"), st.session_state.mistral_api_key, st.session_state.get("previous_context", ""), allow_circular)
                            if extracted_info:
                                extracted_info = smart_historical_data_filling(extracted_info, meeting_date.strftime("%d/%m/%Y"), allow_circular=allow_circular)
                                st.session_state.extracted_info = extracted_info
                                st.success("✅ Information extraction completed!")
                                with st.expander("📋 View Extracted Information"):
                                    st.json(extracted_info)
        else:
            st.info("💡 **Teams Transcript?** Consider uploading the video file for better quality!")
            transcription_input = st.text_area("Entrez la transcription :", height=200, help="Tip: Add speaker names like 'Grace: ...' for better extraction")
            if st.button("Soumettre la Transcription") and transcription_input:
                st.session_state.transcription = transcription_input
                with st.spinner("Extraction des informations avec contexte historique..."):
                    allow_circular = getattr(st.session_state, 'test_mode', False)
                    if allow_circular:
                        st.info("🧪 Test mode: Using existing JSON data for perfect generation")
                    
                    extracted_info = extract_info(st.session_state.transcription, meeting_title, meeting_date.strftime("%d/%m/%Y"), st.session_state.mistral_api_key, st.session_state.get("previous_context", ""), allow_circular)
                    if extracted_info:
                        extracted_info = smart_historical_data_filling(extracted_info, meeting_date.strftime("%d/%m/%Y"), allow_circular=allow_circular)
                        st.session_state.extracted_info = extracted_info
                        st.success("✅ Information extraction completed!")
                        with st.expander("📋 View Extracted Information"):
                            st.json(extracted_info)
        
        if 'extracted_info' in st.session_state and st.button("Générer et Télécharger le Document"):
            with st.spinner("Génération du document avec données complètes..."):
                docx_data = fill_template_and_generate_docx(st.session_state.extracted_info, meeting_title, meeting_date)
                if docx_data:
                    st.success("✅ Document generated successfully!")
                    st.download_button(
                        label="Télécharger le Document",
                        data=docx_data,
                        file_name=f"{meeting_title}_{meeting_date.strftime('%Y-%m-%d')}_notes.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

if __name__ == "__main__":
    main()