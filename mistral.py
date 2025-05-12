import streamlit as st
from mistralai import Mistral
import os
import tempfile
from datetime import datetime
import requests
from transformers import pipeline
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import warnings
import torch
import torchaudio
import json
import re
import base64

# Suppress warnings for cleaner output
warnings.filterwarnings("ignore")

st.set_page_config(page_title="Meeting Transcription Tool", page_icon=":microphone:", layout="wide")

def transcribe_audio(audio_file, file_extension, model_size="base"):
    """Transcribe the uploaded audio file to text using the Whisper model"""
    try:
        model_id_mapping = {
            "tiny": "openai/whisper-tiny",
            "base": "openai/whisper-base",
            "small": "openai/whisper-small",
            "medium": "openai/whisper-medium",
            "large": "openai/whisper-large-v3",
        }
        model_id = model_id_mapping.get(model_size, "openai/whisper-base")
        transcriber = pipeline("automatic-speech-recognition", model=model_id)
        
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_extension}") as temp_audio:
            temp_audio.write(audio_file.getvalue())
            temp_audio_path = temp_audio.name
        
        try:
            waveform, sample_rate = torchaudio.load(temp_audio_path, backend="ffmpeg")
            if sample_rate != 16000:
                waveform = torchaudio.functional.resample(waveform, sample_rate, 16000)
            if waveform.shape[0] > 1:
                waveform = torch.mean(waveform, dim=0, keepdim=True)
            result = transcriber({"raw": waveform[0].numpy(), "sampling_rate": 16000}, chunk_length_s=30, stride_length_s=5)
            return result["text"]
        finally:
            os.unlink(temp_audio_path)
    except Exception as e:
        st.error(f"Error during audio transcription: {e}")
        return f"Error during audio transcription: {e}"

def extract_context_from_report(file, mistral_api_key):
    """Extract text from the uploaded file using Mistral OCR."""
    if not file or not mistral_api_key:
        return ""
    
    file_extension = file.name.split('.')[-1].lower()
    valid_extensions = ['pdf', 'png', 'jpg', 'jpeg']
    if file_extension not in valid_extensions:
        st.error("Unsupported file type.")
        return ""
    
    try:
        client = Mistral(api_key=mistral_api_key)
        
        # Upload the file to Mistral's servers
        uploaded_file = client.files.upload(
            file={
                "file_name": file.name,
                "content": file.getvalue(),
            },
            purpose="ocr"
        )
        
        # Retrieve the signed URL
        signed_url = client.files.get_signed_url(file_id=uploaded_file.id)
        
        # Determine document type
        document_type = "document_url" if file_extension == 'pdf' else "image_url"
        
        # Prepare OCR request
        messages = [
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": "Extract all text from the document."
                    },
                    {
                        "type": document_type,
                        document_type: signed_url.url
                    }
                ]
            }
        ]
        
        # Call Mistral OCR via chat API
        response = client.chat.complete(
            model="mistral-small-latest",
            messages=messages
        )
        
        # Return extracted text
        return response.choices[0].message.content.strip()
    
    except Exception as e:
        st.error(f"Error processing file with Mistral OCR: {e}")
        return ""

def extract_info(transcription, meeting_title, date, deepseek_api_key, previous_context=""):
    """Extract key information from the transcription using Deepseek API with previous context."""
    prompt = f"""
    You are an AI assistant specialized in drafting meeting reports. 
    From the following transcription and the previous meeting context (specifically Activity Review, Resolutions Summary), extract key information and return it as structured JSON in English.

    **Previous Meeting Context**:
    {previous_context if previous_context else "No context available."}

    **Current Meeting Transcription**:
    {transcription}

    **Sections to Extract**:
    - **presence_list**: List of present and absent participants as a string (e.g., "Present: Alice, Bob\nAbsent: Charlie"). Identify names mentioned as present or absent. If not found, use "Present: Not specified\nAbsent: Not specified".
    - **agenda_items**: List of agenda items as a string (e.g., "1. Review of minutes\n2. Resolutions"). Deduce discussed or explicitly mentioned items. If not found, use "Not specified".
    - **resolutions_summary**: List of resolutions as an array (list of dictionaries with keys "date", "dossier", "resolution", "responsible", "deadline", "execution_date", "status", "report_count"). "date", "deadline", and "execution_date" in DD/MM/YYYY format. "report_count" as a string (e.g., "0").
    - **sanctions_summary**: List of sanctions as an array (list of dictionaries with keys "name", "reason", "amount", "date", "status"). "date" in DD/MM/YYYY format, "amount" as a string.
    - **start_time**: Meeting start time (format HHhMMmin, e.g., 07h00min). Deduce if mentioned, else use "Not specified".
    - **end_time**: Meeting end time (format HHhMMmin, e.g., 10h34min). Deduce if mentioned, else use "Not specified".
    - **rapporteur**: Name of the meeting rapporteur. Deduce if mentioned, else use "Not specified".
    - **president**: Name of the meeting president. Deduce if mentioned, else use "Not specified".
    - **balance_amount**: DRI Solidarity account balance (e.g., "827540"). Deduce if mentioned, else use "Not specified".
    - **balance_date**: Balance date (format DD/MM/YYYY). Deduce if mentioned, else use provided date: {date}.

    **Instructions**:
    1. For each speaker, identify their dossier(s), dates, resolutions, responsible party, deadline, status, and report count.
    2. If information is missing, use reasonable defaults (e.g., "Not specified" or provided date: {date}).
    3. Ensure JSON is well-formed and dates follow DD/MM/YYYY format.

    Return the result as structured JSON in English.
    """
    try:
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {deepseek_api_key}"
        }
        payload = {
            "model": "deepseek-chat",
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.1,
            "max_tokens": 4000
        }
        response = requests.post(
            "https://api.deepseek.com/v1/chat/completions",
            headers=headers,
            json=payload
        )
        if response.status_code == 200:
            raw_response = response.json()["choices"][0]["message"]["content"].strip()
            st.write(f"Raw Deepseek response: {raw_response}")
            try:
                extracted_data = json.loads(raw_response)
                extracted_data["date"] = date
                return extracted_data
            except json.JSONDecodeError as e:
                st.error(f"Error parsing JSON: {e}")
                return None
        else:
            st.error(f"Deepseek API error: Status {response.status_code}, Message: {response.text}")
            return None
    except Exception as e:
        st.error(f"Error extracting information: {e}")
        return None

def fill_template_and_generate_docx(extracted_info):
    """Build the Word document from scratch using python-docx"""
    try:
        doc = Document()
        # Add document content generation logic here (simplified for brevity)
        doc.add_heading("Meeting Notes", 0)
        doc.add_paragraph(f"Title: {extracted_info.get('meeting_title', 'Untitled')}")
        doc.add_paragraph(f"Date: {extracted_info.get('date', 'Not specified')}")
        doc.add_paragraph(f"Presence: {extracted_info.get('presence_list', 'Not specified')}")
        doc.add_paragraph(f"Agenda: {extracted_info.get('agenda_items', 'Not specified')}")
        
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            doc.save(tmp.name)
            with open(tmp.name, "rb") as f:
                docx_data = f.read()
            os.unlink(tmp.name)
        return docx_data
    except Exception as e:
        st.error(f"Error generating Word document: {e}")
        return None

def main():
    st.title("Meeting Transcription Tool")
    
    # Sidebar for API keys and previous report
    st.sidebar.header("Configuration")
    st.session_state.mistral_api_key = st.sidebar.text_input("Mistral API Key", type="password")
    st.session_state.deepseek_api_key = st.sidebar.text_input("Deepseek API Key", type="password")
    
    st.sidebar.header("Previous Context")
    previous_report = st.sidebar.file_uploader("Upload Previous Report (optional)", type=["pdf", "png", "jpg", "jpeg"])
    if previous_report:
        status_text = st.sidebar.empty()
        status_text.text("Extracting context...")
        context = extract_context_from_report(previous_report, st.session_state.mistral_api_key)
        status_text.text("Context extracted successfully!")
        st.session_state.previous_context = context
    else:
        st.session_state.previous_context = ""
    
    # Main app content
    col1, col2 = st.columns(2)
    
    with col1:
        st.header("Meeting Details")
        meeting_title = st.text_input("Meeting Title", value="Meeting")
        meeting_date = st.date_input("Meeting Date", datetime.now())
    
    with col2:
        st.header("Transcription & Output")
        uploaded_file = st.file_uploader("Upload an audio file", type=["mp3", "wav", "m4a", "flac"])
        whisper_model = st.selectbox("Whisper Model", ["tiny", "base", "small", "medium", "large"], index=1)
        
        if uploaded_file is not None:
            file_extension = uploaded_file.name.split('.')[-1].lower()
            if st.button("Transcribe Audio"):
                with st.spinner(f"Transcribing with Whisper {whisper_model}..."):
                    transcription = transcribe_audio(uploaded_file, file_extension, whisper_model)
                    if transcription and not transcription.startswith("Error"):
                        st.session_state.transcription = transcription
                        st.text_area("Transcription", transcription, height=200)
        
        if 'transcription' in st.session_state:
            if st.button("Extract Information"):
                with st.spinner("Extracting information..."):
                    extracted_info = extract_info(
                        st.session_state.transcription,
                        meeting_title,
                        meeting_date.strftime("%d/%m/%Y"),
                        st.session_state.deepseek_api_key,
                        st.session_state.get("previous_context", "")
                    )
                    if extracted_info:
                        st.session_state.extracted_info = extracted_info
                        st.text_area("Extracted Information", json.dumps(extracted_info, indent=2), height=300)
            
            if 'extracted_info' in st.session_state:
                if st.button("Generate Document"):
                    with st.spinner("Generating document..."):
                        docx_data = fill_template_and_generate_docx(st.session_state.extracted_info)
                        if docx_data:
                            st.download_button(
                                label="Download Meeting Notes",
                                data=docx_data,
                                file_name=f"{meeting_title}_{meeting_date.strftime('%Y-%m-%d')}_notes.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )

if __name__ == "__main__":
    main()